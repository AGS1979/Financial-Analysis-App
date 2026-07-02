"""Tariff Impact Tracker.

Extracts tariff-related commentary and its financial impact from earnings-call
transcripts / filings (PDF via PyMuPDF or FMP transcripts) using DeepSeek, and
exports the result as a Word report. API keys and the logo are passed in by the
router in app.py.
"""

import fitz
import json
import os
import pandas as pd
import requests
import streamlit as st

from datetime import datetime, timedelta, timezone
from docx import Document
from io import BytesIO
from llm import llm
from utils.net import http_get


def tariff_impact_tracker_app(DEEPSEEK_API_KEY: str, FMP_API_KEY: str, logo_base64_string: str):
    """
    Encapsulates the updated Tariff Impact Tracker functionality with revised UI and restored downloads.
    """
    st.markdown("### 📈 Tariff Impact Tracker")
    st.markdown("Analyze earnings call transcripts or public filings to extract and summarize mentions of tariffs and their financial impact.")

    # --- HELPER FUNCTIONS TO PREPARE DATA ---
    def prepare_table_data(all_analyses):
        """Prepares dataframes for display and download to ensure consistency."""
        if not all_analyses:
            return None, None, None

        table1_data, table2_data, table3_data = [], [], []

        for company_key, analysis in all_analyses.items():
            if not analysis or not isinstance(analysis, dict):
                continue

            company_display = f"{analysis.get('company_name', 'N/A')} ({analysis.get('ticker', company_key.upper())})"
            
            table1_data.append({
                "Company": company_display,
                "Management Commentary": analysis.get('management_commentary', 'No discussion'),
                "Vulnerability": analysis.get('vulnerability', 'No discussion'),
                "Profitability Impact": analysis.get('profitability_impact', 'No discussion'),
                "Pricing Implication": analysis.get('pricing_implication', 'No discussion'),
            })
            table2_data.append({
                "Company": company_display,
                "Demand Sensitivity": analysis.get('demand_sensitivity', 'No discussion'),
                "Guidance Implications": analysis.get('guidance_implications', 'No discussion'),
                "Mitigation Strategies": analysis.get('mitigation_strategies', 'No discussion'),
            })
            table3_data.append({
                "Company": company_display,
                "The Known Unknowns": analysis.get('the_known_unknowns', 'No discussion'),
                "Competitive Positioning": analysis.get('competitive_positioning', 'No discussion'),
            })

        df1 = pd.DataFrame(table1_data) if table1_data else pd.DataFrame()
        df2 = pd.DataFrame(table2_data) if table2_data else pd.DataFrame()
        df3 = pd.DataFrame(table3_data) if table3_data else pd.DataFrame()
        
        return df1, df2, df3

    # --- CORE ANALYSIS LOGIC ---
    @st.cache_data(ttl=3600, show_spinner=False)
    def get_transcript_from_fmp(ticker, year, quarter):
        # This function remains unchanged.
        if not FMP_API_KEY:
            st.error("Error: FMP_API_KEY not found.")
            return None, None
        url = f"https://financialmodelingprep.com/api/v3/earning_call_transcript/{ticker}?quarter={quarter}&year={year}&apikey={FMP_API_KEY}"
        company_profile_url = f"https://financialmodelingprep.com/api/v3/profile/{ticker}?apikey={FMP_API_KEY}"
        try:
            company_name = "N/A"
            profile_response = http_get(company_profile_url)
            profile_response.raise_for_status()
            profile_data = profile_response.json()
            if profile_data and "companyName" in profile_data[0]:
                company_name = profile_data[0]['companyName']
            response = http_get(url)
            response.raise_for_status()
            data = response.json()
            if data and "content" in data[0]:
                return data[0]["content"], company_name
            else:
                st.warning(f"No transcript content found for {ticker} for Q{quarter} {year}.")
                return None, None
        except requests.exceptions.RequestException as e:
            st.error(f"Error fetching data from FMP API for {ticker}: {e}")
            return None, None
        except (IndexError, KeyError):
            st.error(f"Error parsing FMP API response for {ticker}. The data might be empty or in an unexpected format.")
            return None, None

    def extract_text_from_pdf(uploaded_file):
        # This function remains unchanged.
        full_text = ""
        try:
            file_bytes = uploaded_file.getvalue()
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for page in doc:
                    full_text += page.get_text() + "\n"
        except Exception as e:
            st.error(f"An error occurred while reading '{uploaded_file.name}': {e}")
        return full_text

    @st.cache_data(ttl=3600, show_spinner=False)
    def analyze_text_with_deepseek(_text_content, company_name, ticker):
        # MODIFICATION 3: Enhanced prompt for better data capture.
        if not DEEPSEEK_API_KEY:
            st.error("Error: DEEPSEEK_API_KEY not found.")
            return None
        if not _text_content or not _text_content.strip():
            st.warning("Input text is empty. Cannot perform analysis.")
            return None

        prompt = f"""
        As a specialist financial analyst, your task is to meticulously analyze the following corporate document for {company_name} ({ticker}).
        Your entire focus must be on comments related to **tariffs, trade duties, and import taxes**.

        **CRITICAL RULE:** You must extract all specific quantitative data mentioned, such as dollar amounts ($40 million), basis points (170 bps), or percentages (10% to 50%). If specific numbers are mentioned, include them directly in your summary. Do not generalize if specifics are provided. For qualitative points, summarize them concisely. If a topic is not discussed, you MUST return "No discussion".

        Return a single valid JSON object with the following fields:
        - "company_name": "{company_name}"
        - "ticker": "{ticker}"
        - "management_commentary": "A concise summary of the company's overall stance and key messages regarding tariffs."
        - "vulnerability": "Identify the company's financial/operational exposure. Name the specific tariffs (e.g., Section 232), products, and countries involved."
        - "profitability_impact": "How do tariffs affect costs and margins? **Capture all specific financial impacts** (e.g., '$90 million annually', 'reduce operating margins by 170 basis points')."
        - "pricing_implication": "How is the company changing prices due to tariffs? Mention any selective or broad-based price increases."
        - "demand_sensitivity": "How are tariffs expected to impact demand for the company's products? Is the effect positive or negative?"
        - "guidance_implications": "How have tariffs specifically impacted the company's financial guidance or outlook? Mention any quantified impacts (e.g., 'incorporated a 170 basis point tariff impact into Q3 guidance')."
        - "mitigation_strategies": "List the key strategies the company is using to handle tariffs (e.g., supply chain changes, cost savings, negotiations, vertical integration)."
        - "the_known_unknowns": "What are the potential The Known Unknowns, risks, or policy uncertainties mentioned?"
        - "competitive_positioning": "How do tariffs affect the company's competitive position? Do they see it as an advantage or disadvantage?"

        Document Text:
        ---
        {_text_content[:40000]}
        ---
        """
        try:
            content_str = llm.chat(
                [{"role": "user", "content": prompt}],
                temperature=0.1,
                response_format={"type": "json_object"},
                timeout=120,
            )
            return json.loads(content_str)
        except requests.exceptions.RequestException as e:
            st.error(f"Error calling DeepSeek API: {e}")
            return None
        except (json.JSONDecodeError, KeyError) as e:
            st.error(f"Error parsing DeepSeek API JSON response: {e}\nResponse: {content_str}")
            return None

    # --- DISPLAY & DOWNLOAD FUNCTIONS ---
    def display_tariff_tables(df1, df2, df3):
        st.markdown("---")
        st.header("Tariff Impact Analysis")

        # Display Table 1
        st.subheader("Table 1: Overall Impact & Exposure")
        if not df1.empty:
            st.markdown(df1.to_html(escape=False, index=False, justify='left'), unsafe_allow_html=True)
        else:
            st.info("No data available to display.")

        st.markdown("<br>", unsafe_allow_html=True) 

        # Display Table 2
        st.subheader("Table 2: Business & Strategy Impact")
        if not df2.empty:
            st.markdown(df2.to_html(escape=False, index=False, justify='left'), unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)

        # Display Table 3
        st.subheader("Table 3: Future Outlook")
        if not df3.empty:
            st.markdown(df3.to_html(escape=False, index=False, justify='left'), unsafe_allow_html=True)

    def generate_html_report(df1, df2, df3, logo_b64):
        # MODIFICATION 2: New HTML report function for the three-table format.
        styles = """<style>
            body { font-family: 'Poppins', sans-serif; background-color: #f9fafb; padding: 20px; color: #333; }
            h1, h2, h3 { color: #1e1e1e; }
            table { width: 100%; border-collapse: collapse; margin-bottom: 30px; }
            th, td { padding: 12px 15px; text-align: left; border: 1px solid #e0e0e0; vertical-align: top; font-size: 14px; }
            th { background-color: #00416A; color: #ffffff; }
            tr:nth-child(even) { background-color: #f9f9f9; }
            .header { display: flex; justify-content: space-between; align-items: center; padding-bottom: 1rem; border-bottom: 3px solid #00416A; margin-bottom: 2rem; }
            .title { font-size: 2.5rem; font-weight: 700; }
            .logo img { height: 40px; }
            </style>"""
        
        header_html = f"""
            <div class="header">
                <div class="title">Tariff Impact Tracker Report</div>
                <div class="logo"><img src="data:image/png;base64,{logo_b64}" alt="Logo"></div>
            </div>"""

        def df_to_html_bold_company(df):
            df_copy = df.copy()
            df_copy['Company'] = df_copy['Company'].apply(lambda x: f"<b>{x}</b>")
            return df_copy.to_html(escape=False, index=False)

        table1_html = f"<h2>Table 1: Overall Impact & Exposure</h2>" + (df_to_html_bold_company(df1) if not df1.empty else "<p>No data.</p>")
        table2_html = f"<h2>Table 2: Business & Strategy Impact</h2>" + (df_to_html_bold_company(df2) if not df2.empty else "<p>No data.</p>")
        table3_html = f"<h2>Table 3: Future Outlook</h2>" + (df_to_html_bold_company(df3) if not df3.empty else "<p>No data.</p>")
        
        full_html_content = f"<html><head><title>Tariff Impact Report</title>{styles}</head><body>{header_html}{table1_html}{table2_html}{table3_html}</body></html>"
        return full_html_content

    def generate_word_report(df1, df2, df3):
        # MODIFICATION 2: New Word report function for the three-table format.
        doc = Document()
        doc.add_heading('Tariff Impact Report', level=0)
        
        for i, df in enumerate([df1, df2, df3]):
            if df.empty: continue
            
            table_titles = ["Table 1: Overall Impact & Exposure", "Table 2: Business & Strategy Impact", "Table 3: Future Outlook"]
            doc.add_heading(table_titles[i], level=1)
            
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for j, col_name in enumerate(df.columns):
                hdr_cells[j].text = col_name

            for index, row in df.iterrows():
                row_cells = table.add_row().cells
                for j, cell_value in enumerate(row):
                    row_cells[j].text = str(cell_value)
            doc.add_paragraph() # Add space between tables

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # --- STREAMLIT UI LAYOUT ---
    st.subheader("Data Source")
    data_source = st.radio(
        "Choose where to get the transcript from:",
        ("Fetch Transcript", "Upload PDF Transcript(s)"),
        horizontal=True,
        label_visibility="collapsed"
    )

    if 'tariff_all_analysis_results' not in st.session_state:
        st.session_state.tariff_all_analysis_results = {}
    # --- NEW UI for Custom Prompt ---
    st.markdown("---")
    st.subheader("Advanced: Customize Analysis Prompt")
    st.warning("The Tariff Tracker requires a specific JSON output. Your custom prompt must request this structure or the analysis will fail.")
    st.text_area(
        "Enter your custom prompt for tariff analysis:",
        placeholder="Enter your full custom prompt here. It must ask for a JSON object with keys like 'management_commentary', 'vulnerability', etc.",
        height=250,
        key="tariff_custom_prompt"
    )
    # --- END NEW UI ---    
    if data_source == "Fetch Transcript":
        tickers_input = st.text_input("Company Ticker(s)", "CROX, STLD, CLF", help="Enter one or more tickers, separated by commas.")
        c2, c3 = st.columns(2)
        with c2: year = st.number_input("Year", min_value=2010, max_value=datetime.now().year + 1, value=2025)
        with c3: quarter = st.selectbox("Quarter", [1, 2, 3, 4], index=1)

        if st.button("Fetch & Analyze Transcripts", type="primary"):
            tickers = [ticker.strip().upper() for ticker in tickers_input.split(',') if ticker.strip()]
            if tickers:
                st.session_state.tariff_all_analysis_results = {}
                # MODIFICATION 1: Simplified loading message.
                with st.spinner("Generating analysis... This may take a moment."):
                    results = {}
                    for ticker in tickers:
                        text_to_analyze, company_name = get_transcript_from_fmp(ticker, year, quarter)
                        if text_to_analyze:
                            results[ticker] = analyze_text_with_deepseek(text_to_analyze, company_name, ticker)
                    st.session_state.tariff_all_analysis_results = results

    elif data_source == "Upload PDF Transcript(s)":
        uploaded_files = st.file_uploader("Upload one or more PDF files", type="pdf", accept_multiple_files=True)

        if st.button("Upload & Analyze PDFs", type="primary"):
            if uploaded_files:
                st.session_state.tariff_all_analysis_results = {}
                # MODIFICATION 1: Simplified loading message.
                with st.spinner("Generating analysis... This may take a moment."):
                    results = {}
                    for uploaded_file in uploaded_files:
                        company_key = os.path.splitext(uploaded_file.name)[0]
                        text_to_analyze = extract_text_from_pdf(uploaded_file)
                        if text_to_analyze:
                            results[company_key] = analyze_text_with_deepseek(text_to_analyze, company_key, "N/A")
                    st.session_state.tariff_all_analysis_results = results
            else:
                st.warning("Please upload at least one PDF file.")

    # --- DISPLAY RESULTS AND DOWNLOADS ---
    if st.session_state.get('tariff_all_analysis_results'):
        all_results = st.session_state.tariff_all_analysis_results
        df1, df2, df3 = prepare_table_data(all_results)
        
        display_tariff_tables(df1, df2, df3)

        # MODIFICATION 2: Re-enabled download buttons with updated functions.
        st.markdown("---")
        st.header("Download Report")
        
        col1, col2 = st.columns(2)
        if not df1.empty or not df2.empty or not df3.empty:
            with col1:
                html_content = generate_html_report(df1, df2, df3, logo_base64_string)
                st.download_button(
                    label="📥 Download as HTML",
                    data=html_content,
                    file_name="tariff_impact_report.html",
                    mime="text/html"
                )
            with col2:
                word_buffer = generate_word_report(df1, df2, df3)
                st.download_button(
                    label="📥 Download as Word",
                    data=word_buffer,
                    file_name="tariff_impact_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
