"""IPO Investment Memo Generator (Agent Pre-IPO).

Builds a DRHP/IPO investment memo + infographic and answers questions over the
uploaded document via an in-memory FAISS + SentenceTransformer + DeepSeek
retrieval engine (DocumentQueryEngine).
"""

import faiss
import fitz
import markdown
import numpy as np
import os
import re
import requests
import streamlit as st
import tempfile

from PyPDF2 import PdfReader
from bs4 import BeautifulSoup
from collections import defaultdict
from datetime import datetime, timedelta, timezone
from docx import Document
from docx.shared import Pt, Inches
from jinja2 import Template
from pathlib import Path
from sentence_transformers import SentenceTransformer
from utils.logging import log_audit_event, log_user_history, get_user_history
from utils.net import http_post, http_get


def investment_memo_app():
    """
    Encapsulates the entire IPO Investment Memo Generator with Infographic and Q&A.
    This version is aligned with the advanced standalone module and supports both PDF and HTML.
    """
    
    # --- CONFIGURATION ---
    DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY")
    DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
    CHUNK_SIZE = 50

    if not DEEPSEEK_API_KEY:
        st.error("DeepSeek API key not found. Please add it to your Streamlit secrets.")
        return

    # --- HELPER FUNCTIONS (Memo & Pipeline) ---

    def clean_markdown(text):
        text = re.sub(r'#+\s*', '', text)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'_+', '', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'^[-*•]+\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'Section\s\d+[:.]?', '', text, flags=re.IGNORECASE)
        text = re.sub(r'(Next|Previous) section:.*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'This section .*?(focuses on|explores|explains).*?\.', '', text, flags=re.IGNORECASE)
        return text.strip()

    def extract_text_by_page(pdf_path):
        doc = fitz.open(pdf_path)
        return [page.get_text() for page in doc], len(doc)

    # MODIFICATION: New function to handle HTML files
    def extract_text_from_html(html_path):
        """Reads an HTML file and extracts all text content."""
        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
        
        for element in soup(["script", "style", "header", "footer", "nav"]):
            element.decompose() # Remove irrelevant tags
            
        text = soup.get_text(separator='\n', strip=True)
        # Consolidate whitespace
        cleaned_text = re.sub(r'\n{3,}', '\n\n', text)
        return cleaned_text

    def get_relevant_pages_chunked(text_by_page, user_query):
        total_pages = len(text_by_page)
        relevant_pages = set()
        for start in range(0, total_pages, CHUNK_SIZE):
            end = min(start + CHUNK_SIZE, total_pages)
            chunk_pages = text_by_page[start:end]
            prompt = (
                "Below are texts from a PDF. Identify only the page numbers (starting from 1) relevant to this query:\n"
                f"Query: {user_query}\n\n"
            )
            for i, text in enumerate(chunk_pages):
                snippet = text[:1000].replace('\n', ' ')
                prompt += f"\nPage {start + i + 1}: {snippet}\n"
            
            messages = [
                {"role": "system", "content": "You are an expert document analyst."},
                {"role": "user", "content": prompt}
            ]
            response = http_post(DEEPSEEK_API_URL, headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}, json={"model": "deepseek-chat", "messages": messages})
            response.raise_for_status()
            reply = response.json()['choices'][0]['message']['content']
            matches = re.findall(r'\d+', reply)
            for m in matches:
                if 1 <= int(m) <= total_pages:
                    relevant_pages.add(int(m))
        return sorted(relevant_pages)

    def extract_selected_pages_text(original_path, pages_to_keep):
        doc = fitz.open(original_path)
        return "\n".join(doc[p - 1].get_text() for p in pages_to_keep).strip()

    def extract_company_name(text):
        prompt = (
            "Extract only the legal name of the company from the following IPO or DRHP text. "
            "Return only the company name, nothing else.\n\n"
            f"{text[:10000]}"
        )
        messages = [
            {"role": "system", "content": "You are an expert in IPO documents."},
            {"role": "user", "content": prompt}
        ]
        response = http_post(DEEPSEEK_API_URL, headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}, json={"model": "deepseek-chat", "messages": messages})
        response.raise_for_status()
        
        # --- THE FIX IS HERE ---
        # Get the raw response and clean it by replacing newlines with spaces before returning.
        company_name = response.json()['choices'][0]['message']['content'].strip()
        return company_name.replace('\n', ' ')

    def find_relevant_text_for_section(full_text, section_title, keywords_map):
        """
        Searches the full document text for keywords related to a specific section title
        and returns a relevant chunk of text for the LLM to use as context.
        """
        # Get the specific keywords for the given section title
        keywords = keywords_map.get(section_title, [])
        if not keywords:
            # Fallback to a generic chunk if no keywords are defined for a title
            return full_text[:16000]

        # Create a regex pattern to find any of the keywords, ignoring case
        # This pattern looks for the keyword as a potential heading.
        pattern = re.compile(r'^\s*(' + '|'.join(map(re.escape, keywords)) + r')\s*$', re.IGNORECASE | re.MULTILINE)
        
        match = pattern.search(full_text)
        
        if match:
            # If a keyword is found, extract the text following it
            start_index = match.end()
            # Extract a substantial chunk of text to provide enough context
            context_text = full_text[start_index : start_index + 20000]
            return context_text
        else:
            # If no specific section is found, return the initial part of the document as a fallback.
            # This is not ideal but prevents errors. The key is a good keywords_map.
            st.warning(f"Could not find a specific section for '{section_title}'. Using initial text as fallback.")
            return full_text[:16000]

    def generate_memo_sections(filtered_text, custom_notes=""):
        section_titles = [
            "1. IPO Offer Details", "2. Company Overview", "3. Industry Overview and Outlook",
            "4. Business Model", "5. Financial Highlights",
            "6. Guidance and Outlook on future financial performance",
            "7. Peer Comparison and Competitors", "8. Risks", "9. Investment Highlights"
        ]

        # A map to find the right DRHP sections for each memo title. This is crucial.
        keywords_map = {
            "1. IPO Offer Details": ["Details of the Offer", "The Offer"],
            "2. Company Overview": ["Our Business", "Company Overview", "Summary of Business"],
            "3. Industry Overview and Outlook": ["Industry Overview"],
            "4. Business Model": ["Business Model", "Our Business"],
            "5. Financial Highlights": ["Financial Highlights", "Financial Information", "Restated Consolidated Financial Information", "Summary of Restated Consolidated Financial Information"],
            "6. Guidance and Outlook on future financial performance": ["Management's Discussion and Analysis", "Guidance and Outlook"],
            "7. Peer Comparison and Competitors": ["Peer Comparison", "Competitive Landscape", "Competitors"],
            "8. Risks": ["Risk Factors", "Risks"],
            "9. Investment Highlights": ["Investment Rationale", "Investment Highlights", "Our Competitive Strengths"]
        }
        
        sections = {}
        st.info("Generating memo sections with targeted context analysis...")
        progress_bar = st.progress(0)
        
        for i, title in enumerate(section_titles):
            st.write(f"Finding context for: **{title[3:]}**...")
            
            # --- THIS IS THE CORE FIX ---
            # Find relevant text specifically for THIS section, instead of using the same generic text for all.
            relevant_text = find_relevant_text_for_section(filtered_text, title, keywords_map)
            
            st.write(f"Generating content for: **{title[3:]}**...")
            
            prompt = (
                f"You are a professional financial analyst writing a Pre-IPO memo section titled: '{title[3:]}'.\n"
                "Based STRICTLY on the 'Relevant DRHP Text' provided below, generate approximately 500 words of clean, structured, analytical prose suitable for institutional investors.\n"
                "IMPORTANT RULES:\n"
                "- ONLY use information present in the provided text. Do not invent or hallucinate any data, figures, or facts.\n"
                "- If specific information (e.g., financial numbers, competitor names) is not in the text, state that the information is not available in the provided context.\n"
                "- Do not mention this is a memo. Do not start with the section title (e.g., avoid 'The Business Model is...').\n"
                "- Avoid phrases like 'In this section' or 'The provided text states...'. Write authoritatively.\n"
                "- Use plain, professional text only. Strictly avoid markdown (no asterisks, hashes, underscores).\n\n"
            )
            if custom_notes:
                prompt += f"USER'S CUSTOM FOCUS: {custom_notes.strip()}\n\n"
            
            prompt += f"Relevant DRHP Text:\n{relevant_text}" # We use the specifically found relevant_text
            
            messages = [
                {"role": "system", "content": "You are an expert financial analyst writing an investment memo based *only* on provided text."},
                {"role": "user", "content": prompt}
            ]
            
            try:
                response = http_post(
                    DEEPSEEK_API_URL, 
                    headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}, 
                    json={"model": "deepseek-chat", "messages": messages, "temperature": 0.2}
                )
                response.raise_for_status()
                raw_content = response.json()['choices'][0]['message']['content']
                
                # Clean the generated content
                cleaned = clean_markdown(raw_content)
                cleaned = re.sub(rf"^{re.escape(title[3:])}[\s:—-]*", "", cleaned, flags=re.IGNORECASE | re.MULTILINE)
                sections[title] = cleaned.strip()
                
            except requests.RequestException as e:
                st.error(f"API call failed for section '{title[3:]}': {e}")
                sections[title] = f"Error generating this section. {e}"

            progress_bar.progress((i + 1) / len(section_titles))
            
        st.success("All memo sections generated.")
        return sections

    def save_sections_to_word(sections_dict, company_name="Company", output_dir="documents"):
        os.makedirs(output_dir, exist_ok=True)
        
        # Sanitize the company name for the filename
        sanitized_company_name = re.sub(r'[\\/*?:"<>|]', "", company_name)
        sanitized_company_name = sanitized_company_name.replace(' ', '_')
        
        # Truncate to a reasonable length to avoid "File name too long" errors
        max_len = 50 
        if len(sanitized_company_name) > max_len:
            sanitized_company_name = sanitized_company_name[:max_len]

        filename = f"{sanitized_company_name}_PreIPO_Memo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        full_path = os.path.join(output_dir, filename)
        
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Aptos Display'
        style.font.size = Pt(11)

        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"{company_name} Pre-IPO Memo")
        title_run.font.name = 'Aptos Display'
        title_run.font.size = Pt(20)
        title_run.bold = True
        doc.add_paragraph()

        for title, body in sections_dict.items():
            heading = doc.add_paragraph()
            run = heading.add_run(title)
            run.bold = True
            run.font.name = 'Aptos Display'
            run.font.size = Pt(14)
            for para in body.strip().split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.add_paragraph()
        
        section = doc.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        
        doc.save(full_path)
        return full_path

    # MODIFICATION: The run_memo_pipeline is modified below in the main app flow
    # to handle the logic split between PDF and HTML.

    # --- HELPER FUNCTIONS (Infographic) ---

    def extract_raw_text_from_docx(docx_path):
        doc = Document(docx_path)
        return "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())

    def call_deepseek_summary(text, company_name):
        prompt = f"""
        You are an investment analyst. Summarize the key points for each section of the provided pre-IPO memo for {company_name}.
        For each section, provide 3-5 crisp bullet points. Each bullet point must be under 30 words.
        Format your response in markdown, with each section as a header.
        
        Memo Text:
        {text}
        """
        messages = [
            {"role": "system", "content": "You are a financial analyst specializing in concise summaries for infographics."},
            {"role": "user", "content": prompt}
        ]
        response = http_post(DEEPSEEK_API_URL, headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}, json={"model": "deepseek-chat", "messages": messages, "temperature": 0.3})
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    
    def parse_deepseek_response(summary_text):
        sections = defaultdict(list)
        current_section = None
        header_pattern = re.compile(r"^#+\s*(?:\d+\.\s*)?(.*?)\s*$", re.MULTILINE)
        
        parts = header_pattern.split(summary_text)
        if len(parts) > 1:
            for i in range(1, len(parts), 2):
                header = parts[i].strip()
                content = parts[i+1]
                bullets = re.findall(r'[-\*•]\s+(.*)', content)
                formatted_bullets = [re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', b.strip()) for b in bullets]
                if header and formatted_bullets:
                    sections[header] = formatted_bullets
        return dict(sections)

    def generate_infographic_html(docx_path, company_name, template_path="base_infographic.html"):
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"The template file '{template_path}' was not found.")
        raw_text = extract_raw_text_from_docx(docx_path)
        summary = call_deepseek_summary(raw_text, company_name)
        sections = parse_deepseek_response(summary)
        
        with open(template_path, "r", encoding="utf-8") as f:
            html_template = f.read()
            
        template = Template(html_template)
        return template.render(company_name=company_name, sections=sections)

    # --- HELPER FUNCTIONS (Q&A Engine) ---

    # MODIFICATION: Renamed and updated to handle both PDF and HTML
    class DocumentQueryEngine:
        def __init__(self, model_name="all-MiniLM-L6-v2"):
            self.api_key = DEEPSEEK_API_KEY
            self.embedder = SentenceTransformer(model_name)

        def _extract_chunks_from_pdf(self, path):
            reader = PdfReader(path)
            return [(i + 1, page.extract_text().strip()) for i, page in enumerate(reader.pages) if page.extract_text()]

        def _extract_chunks_from_html(self, path):
            full_text = extract_text_from_html(path)
            # Chunk HTML text by paragraphs (or a fixed character length if needed)
            paragraphs = [p.strip() for p in full_text.split('\n\n') if p.strip()]
            # Return with a chunk index instead of a page number
            return [(i + 1, para) for i, para in enumerate(paragraphs)]

        def answer_query(self, doc_path, query, top_k=3):
            # MODIFICATION: Select the correct chunking method based on file type
            if doc_path.endswith('.pdf'):
                chunks = self._extract_chunks_from_pdf(doc_path)
                source_label = "Page"
            elif doc_path.endswith('.html'):
                chunks = self._extract_chunks_from_html(doc_path)
                source_label = "Section" # Use a more generic label for HTML chunks
            else:
                raise ValueError("Unsupported file type for Q&A.")

            if not chunks: raise ValueError("No text could be extracted from the document.")
            
            source_ids, texts = zip(*chunks)
            text_embeddings = np.array(self.embedder.encode(texts, convert_to_numpy=True))
            
            index = faiss.IndexFlatL2(text_embeddings.shape[1])
            index.add(text_embeddings)
            
            query_embedding = self.embedder.encode([query])
            _, I = index.search(query_embedding, k=top_k)
            
            context_chunks = [(source_ids[i], texts[i]) for i in I[0]]
            
            messages = [{"role": "system", "content": "Answer the user's question based on the context provided from the document."}]
            for source_id, text in context_chunks:
                messages.append({"role": "user", "content": f"[Context from {source_label} {source_id}]:\n{text}"})
            messages.append({"role": "user", "content": f"Question: {query}"})
            
            response = http_post(DEEPSEEK_API_URL, headers={"Authorization": f"Bearer {self.api_key}"}, json={"model": "deepseek-chat", "messages": messages, "temperature": 0.2})
            response.raise_for_status()
            
            answer_md = response.json()["choices"][0]["message"]["content"]
            cited_sources = sorted([source_ids[i] for i in I[0]])
            return markdown.markdown(answer_md), cited_sources, source_label

    # Initialize session state
    if "memo_generated" not in st.session_state:
        st.session_state.memo_generated = False
    if "memo_path" not in st.session_state:
        st.session_state.memo_path = None
    if "doc_path" not in st.session_state: # MODIFICATION: Renamed for generality
        st.session_state.doc_path = None
    
    # --- Main App Flow ---
    st.markdown("### 📝 Agent Pre-IPO")
    st.markdown("Upload a DRHP or IPO prospectus to automatically generate a detailed investment memo, an infographic, and perform Q&A.")
    st.subheader("📤 1. Upload DRHP or IPO Prospectus")
    
    # MODIFICATION: Allow both PDF and HTML files
    uploaded_file = st.file_uploader(
        "Upload your PDF or HTML file",
        type=["pdf", "html"],
        key="prospectus_uploader"
    )
    
    if uploaded_file:
        # Save uploaded file to a temporary path with the correct extension
        suffix = Path(uploaded_file.name).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(uploaded_file.getbuffer())
            st.session_state.doc_path = tmp_file.name
            
        custom_focus = st.text_area(
            "Optional: Add custom notes to guide memo generation",
            key="memo_focus",
            help="Example: 'Focus on the competitive landscape in North America and risks related to supply chain.'"
        )

        # --- NEW UI for Custom Section Prompts ---
        st.markdown("---")
        st.subheader("Advanced: Customize Section Prompts")
        st.info("You can provide your own generation prompt for each memo section below. Leave a box empty to use the default.")
        
        section_titles = [
            "1. IPO Offer Details", "2. Company Overview", "3. Industry Overview and Outlook",
            "4. Business Model", "5. Financial Highlights",
            "6. Guidance and Outlook on future financial performance",
            "7. Peer Comparison and Competitors", "8. Risks", "9. Investment Highlights"
        ]

        for title in section_titles:
            st.markdown(f"##### Custom Prompt for: {title[3:]}")
            st.text_area(
                label=f"Custom prompt for {title[3:]}",
                placeholder=f"Enter your full custom prompt for the '{title[3:]}' section here...",
                height=150,
                key=f"memo_custom_prompt_{title.replace(' ', '_')}", # Unique key
                label_visibility="collapsed"
            )
        # --- END NEW UI ---

        # In your main app flow, under the "if st.button..."
        if st.button("📘 Generate Investment Memo", key="gen_memo"):
            # --- ADD AUDIT LOG CALL ---
            log_audit_event(
                action_type="PRE_IPO_MEMO_GEN",
                status="STARTED",
                target_id=uploaded_file.name,
                details={"custom_focus": custom_focus}
            )
            # ---
            with st.spinner("⏳ Analyzing document and generating memo... This may take a few minutes."):
                try:
                    full_text = ""
                    if st.session_state.doc_path.endswith(".pdf"):
                        # The extract_text_by_page function returns a list of strings (one per page)
                        text_by_page, _ = extract_text_by_page(st.session_state.doc_path)
                        # Join them all into one single large string
                        full_text = "\n".join(text_by_page)
                    
                    elif st.session_state.doc_path.endswith(".html"):
                        full_text = extract_text_from_html(st.session_state.doc_path)

                    if not full_text.strip():
                        raise ValueError("Could not extract text from the document.")
                    
                    company_name = extract_company_name(full_text)
                    st.session_state.pre_ipo_company_name = company_name # Save for later
                    
                    # --- CALL THE NEW, FIXED FUNCTION ---
                    # It now takes the full_text and handles context internally.
                    sections_dict = generate_memo_sections(full_text, custom_focus)
                    
                    memo_path = save_sections_to_word(sections_dict, company_name=company_name)

                    st.session_state.memo_generated = True
                    st.session_state.memo_path = memo_path
                    st.success("✅ Memo generated successfully!")
                    # --- ADD AUDIT LOG CALL ---
                    log_audit_event(action_type="PRE_IPO_MEMO_GEN", status="SUCCESS", target_id=company_name)
                    # ---
                    
                    # --- START: NEW HISTORY LOG CALL ---
                    log_user_history(
                        action_type="Pre-IPO Memo",
                        target_id=company_name,
                        summary=f"Generated Pre-IPO Memo for {company_name}",
                        details={"source_file": uploaded_file.name}
                    )
                    # --- END: NEW HISTORY LOG CALL ---


                except Exception as e:
                    # --- ADD AUDIT LOG CALL ---
                    log_audit_event(
                        action_type="PRE_IPO_MEMO_GEN",
                        status="FAILURE",
                        target_id=uploaded_file.name,
                        details={"error": str(e)}
                    )
                    # ---
                    st.error(f"❌ Error generating memo: {e}")
                    st.session_state.memo_generated = False
        
        # --- Post-Generation Options ---
        if st.session_state.memo_generated and st.session_state.memo_path:
            with open(st.session_state.memo_path, "rb") as f:
                st.download_button(
                    "📥 Download Memo (.docx)", f, Path(st.session_state.memo_path).name,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            st.markdown("---")
            st.subheader("🎨 2. Generate Infographic")
            if st.button("🖼️ Generate Infographic", key="gen_infographic"):
                company_name = st.session_state.get("pre_ipo_company_name", "Company")
                # --- ADD AUDIT LOG CALL ---
                log_audit_event(action_type="PRE_IPO_INFOGRAPHIC_GEN", status="STARTED", target_id=company_name)
                # ---
                with st.spinner("✨ Creating infographic summary..."):
                    try:
                        infographic_html = generate_infographic_html(st.session_state.memo_path, company_name)
                        
                        st.components.v1.html(infographic_html, width=1100, height=1000, scrolling=True)
                        
                        # --- ADD AUDIT LOG CALL ---
                        log_audit_event(action_type="PRE_IPO_INFOGRAPHIC_GEN", status="SUCCESS", target_id=company_name)
                        # ---
                                        
                        st.download_button(
                            label="📥 Download Infographic (.html)",
                            data=infographic_html,
                            file_name=f"{company_name.replace(' ', '_')}_Infographic.html",
                            mime="text/html"
                        )
                    except Exception as e:
                        # --- ADD AUDIT LOG CALL ---
                        log_audit_event(
                            action_type="PRE_IPO_INFOGRAPHIC_GEN",
                            status="FAILURE",
                            target_id=company_name,
                            details={"error": str(e)}
                        )
                        # ---
                        st.error(f"❌ Error generating infographic: {e}")

        # --- Q&A Section ---
        st.markdown("---")
        st.subheader("🔍 3. Ask Questions from the Document")
        query = st.text_input("Type your question (e.g., What are the key risk factors?)", key="memo_query")
        if query:
            # --- ADD AUDIT LOG CALL ---
            log_audit_event(
                action_type="PRE_IPO_QA",
                status="STARTED",
                target_id=st.session_state.doc_path,
                details={"query": query}
            )
            # ---
            with st.spinner("💬 Searching for answers in the document..."):
                try:
                    # MODIFICATION: Use the updated engine
                    engine = DocumentQueryEngine()
                    answer_html, cited_sources, source_label = engine.answer_query(st.session_state.doc_path, query)
                    st.markdown(answer_html, unsafe_allow_html=True)
                    st.caption(f"📄 Answer generated from information on {source_label.lower()}s: {', '.join(map(str, cited_sources))}")
                    # --- ADD AUDIT LOG CALL ---
                    log_audit_event(
                        action_type="PRE_IPO_QA",
                        status="SUCCESS",
                        target_id=st.session_state.doc_path,
                        details={"query": query}
                    )
                    # ---
                except Exception as e:
                    # --- ADD AUDIT LOG CALL ---
                    log_audit_event(
                        action_type="PRE_IPO_QA",
                        status="FAILURE",
                        target_id=st.session_state.doc_path,
                        details={"error": str(e)}
                    )
                    # ---
                    st.error(f"❌ Query Error: {e}")
