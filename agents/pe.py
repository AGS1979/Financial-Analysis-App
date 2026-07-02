"""Agent PE.

Analyzes confidential PE deal documents in a secure environment (Azure Document
Intelligence + Azure OpenAI): diligence Q&A, key-term comparison, outreach drafting.
"""

from config import require_env
from llm import llm
from utils.net import http_post, http_get


def pe_agent_app_azure():
    """
    A secure, confidential agent for Private Equity analysis using Azure services.
    This version includes a deal document analyzer, an advanced bulk outreach email generator,
    a Diligence Q&A tool, an Expert Call Summarizer, and a Key Terms Comparison tool.
    """
    # --- Local imports ---
    import io, re, html, os, json, uuid
    import markdown, pdfplumber
    import pandas as pd
    import streamlit as st
    import requests
    from bs4 import BeautifulSoup
    from azure.core.credentials import AzureKeyCredential
    from azure.ai.documentintelligence import DocumentIntelligenceClient
    from azure.ai.documentintelligence.models import ContentFormat
    from openai import AzureOpenAI
    from docx import Document
    from docx.shared import Pt

    st.markdown("### 🔒 Agent PE")
    st.markdown(
        "Analyze deal documents, generate outreach emails, and perform diligence with enterprise-grade privacy."
    )

    # --- AGENT CONFIG ---
    _cfg = require_env(
        "AZURE_DI_ENDPOINT", "AZURE_DI_KEY", "AZURE_OPENAI_ENDPOINT",
        "AZURE_OPENAI_KEY", "AZURE_OPENAI_DEPLOYMENT_NAME",
    )
    di_endpoint = _cfg["AZURE_DI_ENDPOINT"]
    di_key = _cfg["AZURE_DI_KEY"]
    openai_endpoint = _cfg["AZURE_OPENAI_ENDPOINT"]
    openai_key = _cfg["AZURE_OPENAI_KEY"]
    openai_deployment_name = _cfg["AZURE_OPENAI_DEPLOYMENT_NAME"]

    # --- PROMPTS FOR DOCUMENT ANALYSIS (Fully expanded) ---
    ANALYSIS_PROMPTS = {
        "Investment Thesis": (
            "You are a top-tier private equity analyst. Generate a comprehensive investment thesis based on the provided context. "
            "**CRITICAL RULE: Your entire response must be in clean MARKDOWN format.** "
            "Use markdown headings (`## Subheading`) for sections and bullet points (`* Point`) for lists. Do not create a main title. "
            "Structure your response with the following markdown headings:\n"
            "## Market Opportunity\n## Competitive Moat\n## Value Creation Levers\n## Overall Rationale"
        ),
        "Key Risks & Mitigants": (
            "You are a senior risk officer. **CRITICAL RULE: Your entire response must be in clean MARKDOWN format.** "
            "Use markdown headings (`## Subheading`) for sections and bullet points (`* Point`) for lists. Do not create a main title. "
            "Structure your response with the following markdown headings:\n"
            "## Market & Competitive Risks\n## Operational Risks\n## Financial Risks"
        ),
        "Financial Highlights": (
            "You are a financial diligence expert. **CRITICAL RULE: Your entire response must be in clean MARKDOWN format.** "
            "Use markdown headings (`## Subheading`) for sections and bullet points (`* Point`) for lists. Do not create a main title. "
            "Structure your response with the following markdown headings:\n"
            "## Revenue & Profitability\n## Balance Sheet Health\n## Cash Flow"
        ),
        "Potential Exit Options": (
            "You are a partner on the investment committee. **CRITICAL RULE: Your entire response must be in clean MARKDOWN format.** "
            "Use markdown headings (`## Subheading`) for sections and bullet points (`* Point`) for lists. Do not create a main title. "
            "Structure your response with the following markdown headings:\n"
            "## Strategic Sale\n## Secondary Buyout\n## Initial Public Offering (IPO)"
        ),
    }

    # --- HELPER FUNCTIONS ---

    # Document Analysis helpers
    def parse_pdf_with_azure_di(file_bytes: bytes) -> tuple[str, list]:
        try:
            di_client = DocumentIntelligenceClient(
                endpoint=di_endpoint,
                credential=AzureKeyCredential(di_key),
            )
            stream = io.BytesIO(file_bytes)
            poller = di_client.begin_analyze_document(
                model_id="prebuilt-layout",
                analyze_request=stream,
                content_type="application/pdf",
                pages=None,
                output_content_format=ContentFormat.MARKDOWN,
            )
            result = poller.result()
            return (result.content or ""), []
        except Exception as e:
            st.error(f"Azure Document Intelligence error: {e}")
            return None, []

    def fallback_pdf_text(file_bytes: bytes) -> str:
        text = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text).strip()
    
    def parse_text_file(file_bytes: bytes) -> str:
        try:
            return file_bytes.decode('utf-8')
        except Exception as e:
            st.warning(f"Could not read text file: {e}")
            return ""

    def parse_excel_to_markdown(file_bytes: bytes, file_name: str) -> str:
        try:
            xls = pd.ExcelFile(io.BytesIO(file_bytes))
            markdown_texts = []
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                df.dropna(how='all', axis=0, inplace=True)
                df.dropna(how='all', axis=1, inplace=True)
                if not df.empty:
                    markdown_texts.append(f"## Data from Sheet: {sheet_name}\n\n")
                    markdown_texts.append(df.to_markdown(index=False))
            return "\n\n".join(markdown_texts)
        except Exception as e:
            st.warning(f"Could not process Excel file {file_name}: {e}")
            return ""

    def analyze_document_with_azure_openai(_context: str, _prompt: str) -> str:
        try:
            return llm.chat(
                [
                    {"role": "system", "content": "You are an expert financial analyst that responds only with clean, structured markdown as instructed."},
                    {"role": "user", "content": f"CONTEXT DOCUMENT:\n---\n{_context}\n---\nYOUR TASK: {_prompt}"},
                ],
                provider="azure",
                model=openai_deployment_name,
            )
        except Exception as e:
            return f"## Error\n\n**Error during Azure OpenAI analysis:** {e}"

    def parse_markdown_to_html(analysis_results: dict) -> tuple[str, str]:
        styles = """
        <style>
            .analysis-container { font-family: 'Poppins', sans-serif; border: 1px solid #e0e0e0; border-radius: 8px; padding: 25px; background-color: #f9fafb; }
            .analysis-container h2 { font-size: 1.5em; font-weight: 600; color: #00416A; border-bottom: 2px solid #00416A; padding-bottom: 10px; margin-top: 0; margin-bottom: 20px; }
            .analysis-container h3 { font-size: 1.05em; font-weight: 600; color: #00416A; padding-bottom: 5px; margin-top: 25px; border-bottom: 1px solid #e6f1f6;}
            .analysis-container p { margin-bottom: 1em; line-height: 1.6; color: #333; }
            .analysis-container ul { list-style-position: outside; padding-left: 20px; margin-top: 1em; margin-bottom: 1em; }
            .analysis-container li { margin-bottom: 0.75em; line-height: 1.6; }
            .analysis-container table { width: 100%; border-collapse: collapse; margin: 15px 0; }
            .analysis-container th, .analysis-container td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            .analysis-container th { background-color: #f2f2f2; }
        </style>
        """
        full_html_body = ""
        for title, markdown_content in analysis_results.items():
            full_html_body += f"<h2>{html.escape(title)}</h2>"
            html_from_md = markdown.markdown(markdown_content, extensions=['tables'])
            processed_html = re.sub(r"<h2>(.*?)</h2>", r"<h3>\1</h3>", html_from_md)
            full_html_body += processed_html
        content_div = f"<div class='analysis-container'>{full_html_body}</div>"
        return styles, content_div

    # ADVANCED OUTREACH HELPER FUNCTIONS
    def scrape_website_text(url: str) -> str:
        if not (url.startswith('http://') or url.startswith('https://')):
            url = 'http://' + url
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
            response = http_get(url, headers=headers, timeout=15)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')
            for element in soup(["script", "style", "nav", "footer", "header"]):
                element.decompose()
            text = soup.get_text(separator='\n', strip=True)
            return re.sub(r'\n{3,}', '\n\n', text)
        except requests.RequestException as e:
            st.warning(f"Failed to scrape {url}: {e}")
            return ""

    def parse_word_doc(file_bytes: bytes) -> str:
        try:
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception:
            st.warning("Failed to parse Word document. Please ensure it is a .docx file.")
            return ""

    def analyze_source_for_outreach(company_name: str, source_text: str) -> str:
        prompt = f"""...""" # Omitted for brevity
        try:
            return llm.chat([{"role": "user", "content": f"{prompt}\n\nSOURCE TEXT:\n---\n{source_text[:12000]}\n---"}], provider="azure", model=openai_deployment_name, temperature=0.1)
        except Exception as e:
            return f"Error during analysis: {e}"

    def generate_advanced_outreach_email(company_name: str, recipient_name: str, analysis_points: str, value_prop: str, sender_name: str, sender_title: str, firm_name: str) -> str:
        prompt = f"""...""" # Omitted for brevity
        try:
            return llm.chat([{"role": "user", "content": prompt}], provider="azure", model=openai_deployment_name, temperature=0.5)
        except Exception as e:
            return f"Error generating email: {e}"
            
    def generate_word_document_from_drafts(drafts: list) -> bytes:
        doc = Document()
        doc.add_heading('Bulk Outreach Email Drafts', level=0)
        style = doc.styles['Normal']
        style.font.name = 'Aptos Display'
        style.font.size = Pt(11)
        for i, item in enumerate(drafts):
            company_name = item.get('Company', 'Unknown Company').replace('_', ' ')
            email_draft = item.get('Draft', 'No draft generated.')
            doc.add_heading(company_name, level=2)
            doc.add_paragraph(email_draft)
            if i < len(drafts) - 1:
                doc.add_page_break()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # --- UI TABS ---
    tab_titles = [
        "Deal Document Analysis", 
        "Diligence Q&A", 
        "Expert Call Summarizer",
        "Key Terms Comparison",
        "Email Outreach Generation"
    ]
    tab1, tab2, tab3, tab4, tab5 = st.tabs(tab_titles)

    # --- TAB 1: DEAL DOCUMENT ANALYSIS ---
    with tab1:
        st.subheader("Analyze Confidential Deal Documents")
        st.info("Upload a CIM, Teaser, or other confidential documents to generate a structured analysis.")
        
        uploaded_files_t1 = st.file_uploader(
            "Upload Teasers, CIMs, or Financials (PDF, XLSX, XLS)", type=["pdf", "xlsx", "xls"],
            accept_multiple_files=True, key="pe_agent_uploader_azure"
        )
        if uploaded_files_t1 and "pe_agent_text" not in st.session_state:
            if st.button("Process Documents", type="primary", key="process_docs_t1"):
                with st.spinner("Processing documents..."):
                    all_texts = []
                    for doc in uploaded_files_t1:
                        file_bytes = doc.getvalue()
                        
                        st.write(f"Processing '{doc.name}'...")
                        file_ext = os.path.splitext(doc.name)[1].lower()
                        doc_content = ""
                        if file_ext == ".pdf":
                            text, _ = parse_pdf_with_azure_di(file_bytes)
                            if not text:
                                text = fallback_pdf_text(file_bytes)
                            doc_content = text
                        elif file_ext in [".xlsx", ".xls"]:
                            doc_content = parse_excel_to_markdown(file_bytes, doc.name)
                        if doc_content:
                            all_texts.append(f"--- START: {doc.name} ---\n{doc_content}\n--- END: {doc.name} ---")
                    if all_texts:
                        st.session_state.pe_agent_text = "\n\n".join(all_texts)
                        st.rerun()
        if "pe_agent_text" in st.session_state:
            st.success("✅ Documents processed and ready for analysis.")
            analysis_choices = st.multiselect("Choose analyses:", options=list(ANALYSIS_PROMPTS.keys()), default=list(ANALYSIS_PROMPTS.keys()))
            if st.button("Generate Analysis", use_container_width=True, key="generate_analysis_t1"):
                full_text = st.session_state.pe_agent_text
                analysis_results = {}
                with st.spinner("Generating insights..."):
                    for choice in analysis_choices:
                        result = analyze_document_with_azure_openai(full_text, ANALYSIS_PROMPTS[choice])
                        analysis_results[choice] = result
                st.session_state.pe_agent_analysis_results = analysis_results
        if "pe_agent_analysis_results" in st.session_state:
            st.markdown("---")
            st.subheader("Generated Analysis")
            styles_html, content_html = parse_markdown_to_html(st.session_state.pe_agent_analysis_results)
            st.markdown(styles_html, unsafe_allow_html=True)
            st.markdown(content_html, unsafe_allow_html=True)


    # --- TAB 2: DILIGENCE Q&A ---
    with tab2:
        st.subheader("Diligence Q&A Agent")
        st.info("Upload a market study, report, or CIM and ask specific questions to get answers directly from the text.")
        
        uploaded_file_t2 = st.file_uploader("Upload Document for Q&A (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt'], key="qa_uploader")
        
        if uploaded_file_t2:
            file_bytes = uploaded_file_t2.getvalue()
            file_ext = os.path.splitext(uploaded_file_t2.name)[1].lower()
            doc_text = ""
            if file_ext == '.pdf':
                doc_text, _ = parse_pdf_with_azure_di(file_bytes)
                if not doc_text: doc_text = fallback_pdf_text(file_bytes)
            elif file_ext == '.docx':
                doc_text = parse_word_doc(file_bytes)
            elif file_ext == '.txt':
                doc_text = parse_text_file(file_bytes)
            
            st.session_state['qa_doc_text'] = doc_text
            st.success(f"✅ Successfully processed '{uploaded_file_t2.name}'. You can now ask questions.")

        if 'qa_doc_text' in st.session_state:
            user_question = st.text_input("Ask a question about the document:")
            if user_question:
                with st.spinner("Searching for answers in the document..."):
                    prompt = f"""You are a Q&A assistant. Answer the user's question based ONLY on the provided document context. If the answer is not in the text, state that clearly. Provide relevant quotes where possible.

                    DOCUMENT CONTEXT:
                    ---
                    {st.session_state['qa_doc_text']}
                    ---
                    
                    QUESTION: {user_question}
                    """
                    st.markdown(llm.chat([{"role": "user", "content": prompt}], provider="azure", model=openai_deployment_name))

    # --- TAB 3: EXPERT CALL SUMMARIZER ---
    with tab3:
        st.subheader("Expert Call Summarizer")
        st.info("Upload one or more transcripts from expert network calls to generate a structured summary.")
        
        uploaded_files_t3 = st.file_uploader("Upload Transcripts (.txt, .docx)", type=['txt', 'docx'], accept_multiple_files=True, key="expert_call_uploader")
        
        if st.button("Summarize Transcripts", key="summarize_calls", disabled=not uploaded_files_t3):
            all_transcripts = []
            for file in uploaded_files_t3:
                file_bytes = file.getvalue()
                file_ext = os.path.splitext(file.name)[1].lower()
                transcript_text = ""
                if file_ext == '.docx':
                    transcript_text = parse_word_doc(file_bytes)
                elif file_ext == '.txt':
                    transcript_text = parse_text_file(file_bytes)
                
                if transcript_text:
                    all_transcripts.append(f"--- TRANSCRIPT: {file.name} ---\n{transcript_text}\n--- END TRANSCRIPT ---")
            
            if all_transcripts:
                full_context = "\n\n".join(all_transcripts)
                prompt = f"""You are a private equity analyst. Synthesize the provided expert call transcript(s) into a single, structured summary. 

                CONTEXT:
                ---
                {full_context}
                ---

                TASK: Generate a report in MARKDOWN format with the following headings:
                ## Key Takeaways
                (A bulleted list of the most critical insights and conclusions from the calls.)
                ## Red Flags & Concerns
                (A bulleted list of points raised by the expert(s) that require further diligence or present risks.)
                ## Supporting Quotes
                (A bulleted list of the most impactful direct quotes from the expert(s), citing the source transcript name if possible.)
                """
                with st.spinner("Synthesizing expert calls..."):
                    st.session_state['expert_call_summary'] = llm.chat([{"role": "user", "content": prompt}], provider="azure", model=openai_deployment_name)

        if 'expert_call_summary' in st.session_state:
            st.markdown(st.session_state['expert_call_summary'])


    # --- TAB 4: KEY TERMS COMPARISON ---
    with tab4:
        st.subheader("Key Terms Comparison Tool")
        st.info("Upload multiple documents (e.g., term sheet, credit agreement) to extract and compare key terms side-by-side.")
        
        uploaded_files_t4 = st.file_uploader("Upload Documents to Compare (.pdf, .docx)", type=['pdf', 'docx'], accept_multiple_files=True, key="comparison_uploader")
        
        terms_to_extract = st.text_area("Enter key terms to extract (one per line)", "Purchase Price\nClosing Conditions\nBreak Fee\nFinancial Covenants")
        
        if st.button("Extract & Compare Terms", key="compare_terms", disabled=(len(uploaded_files_t4) < 2)):
            comparison_data = {}
            terms_list = [term.strip() for term in terms_to_extract.split('\n') if term.strip()]
            
            with st.spinner("Extracting terms from documents..."):
                for file in uploaded_files_t4:
                    file_bytes = file.getvalue()
                    file_ext = os.path.splitext(file.name)[1].lower()
                    doc_text = ""
                    if file_ext == '.pdf':
                        doc_text, _ = parse_pdf_with_azure_di(file_bytes)
                        if not doc_text: doc_text = fallback_pdf_text(file_bytes)
                    elif file_ext == '.docx':
                        doc_text = parse_word_doc(file_bytes)
                    
                    if doc_text:
                        prompt = f"""You are a legal diligence assistant. From the document text provided, extract the exact clause or definition for each of the following key terms. If a term is not found, state "Not Found".

                        DOCUMENT TEXT:
                        ---
                        {doc_text[:20000]}
                        ---

                        KEY TERMS TO EXTRACT:
                        {', '.join(terms_list)}

                        Return your response as a JSON object where keys are the terms and values are the extracted text.
                        """
                        try:
                            extracted_terms = json.loads(llm.chat([{"role": "user", "content": prompt}], provider="azure", model=openai_deployment_name, response_format={"type": "json_object"}))
                            comparison_data[file.name] = extracted_terms
                        except json.JSONDecodeError:
                            st.warning(f"Could not parse JSON for {file.name}")
            
            if comparison_data:
                df_data = []
                for term in terms_list:
                    row = {"Key Term": term}
                    for doc_name, terms in comparison_data.items():
                        row[doc_name] = terms.get(term, "Not Found")
                    df_data.append(row)
                
                st.session_state['comparison_df'] = pd.DataFrame(df_data)

        if 'comparison_df' in st.session_state:
            st.dataframe(st.session_state['comparison_df'], use_container_width=True)


    # --- TAB 5: ADVANCED OUTREACH GENERATION ---
    with tab5:
        st.subheader("Generate Highly Personalized Outreach Emails")
        st.info(
            "💡 **Pro Tip:** Some websites may block automated access. For reliable results, use the 'Upload Document' option.",
            icon="ℹ️"
        )
        st.markdown("""<style>.section-header { font-weight: 600; font-size: 1rem; padding-bottom: 0; margin-bottom: -0.5rem; }</style>""", unsafe_allow_html=True)
        
        if 'pe_outreach_targets' not in st.session_state:
            st.session_state.pe_outreach_targets = []
        if not st.session_state.pe_outreach_targets:
            new_id = str(uuid.uuid4())
            st.session_state.pe_outreach_targets.append({'id': new_id, 'file_uploader_key': f'file_{new_id}'})
        
        num_targets = st.number_input(
            "Number of Target Companies", min_value=1, max_value=20,
            value=len(st.session_state.pe_outreach_targets), step=1, key='num_outreach_targets'
        )

        current_len = len(st.session_state.pe_outreach_targets)
        if current_len < num_targets:
            for _ in range(num_targets - current_len):
                new_id = str(uuid.uuid4())
                st.session_state.pe_outreach_targets.append({'id': new_id, 'file_uploader_key': f'file_{new_id}'})
        elif current_len > num_targets:
            st.session_state.pe_outreach_targets = st.session_state.pe_outreach_targets[:num_targets]

        with st.form("advanced_outreach_form"):
            st.markdown("<p class='section-header'>1. Define Your Firm & Sender Details</p>", unsafe_allow_html=True)
            firm_value_prop = st.text_area("Your Firm's Value Proposition", placeholder="e.g., We are a growth equity firm...", label_visibility="collapsed")
            c1, c2, c3 = st.columns(3)
            sender_name = c1.text_input("Your Name", placeholder="Alex Johnson")
            sender_title = c2.text_input("Your Title", placeholder="Associate")
            firm_name = c3.text_input("Your Firm's Name", placeholder="Growth Equity Partners")

            st.markdown('<p class="section-header">2. Define Target Companies</p>', unsafe_allow_html=True)
            
            for i, target in enumerate(st.session_state.pe_outreach_targets):
                st.markdown(f"---")
                st.markdown(f"**Target Company #{i+1}**")
                cols = st.columns([4, 4, 5])
                target['company_name'] = cols[0].text_input("Company Name", key=f"company_{target['id']}")
                target['recipient_name'] = cols[1].text_input("Recipient Name (Optional)", key=f"recipient_{target['id']}")
                target['source_type'] = cols[2].radio("Customization Source", ["Website URL", "Upload Document"], key=f"source_{target['id']}", horizontal=True)

                if target['source_type'] == "Website URL":
                    target['url'] = st.text_input("Website URL", placeholder="www.examplecorp.com", key=f"url_{target['id']}")
                else:
                    target['file'] = st.file_uploader("Upload PDF or Word Doc", type=['pdf', 'docx'], key=target['file_uploader_key'])
            
            submitted = st.form_submit_button("🚀 Generate Email Drafts", use_container_width=True)

        if submitted:
            if not all([firm_value_prop, sender_name, firm_name]):
                st.warning("Please fill out your firm and sender details.")
            else:
                email_drafts = []
                with st.spinner("Analyzing sources and generating drafts..."):
                    for i, target_data in enumerate(st.session_state.pe_outreach_targets):
                        company = target_data.get('company_name')
                        if not company:
                            st.warning(f"Skipping Target #{i+1} because company name is missing.")
                            continue

                        st.write(f"Processing: **{company}**")
                        source_text = ""
                        uploaded_file_obj = target_data.get('file')

                        if target_data['source_type'] == "Website URL" and target_data.get('url'):
                            source_text = scrape_website_text(target_data['url'])
                        elif target_data['source_type'] == "Upload Document" and uploaded_file_obj:
                            file_bytes = uploaded_file_obj.getvalue()
                            if uploaded_file_obj.name.lower().endswith('.pdf'):
                                source_text, _ = parse_pdf_with_azure_di(file_bytes)
                                if not source_text: source_text = fallback_pdf_text(file_bytes)
                            elif uploaded_file_obj.name.lower().endswith('.docx'):
                                source_text = parse_word_doc(file_bytes)
                        
                        analysis = analyze_source_for_outreach(company, source_text)
                        email = generate_advanced_outreach_email(
                            company, target_data.get('recipient_name'), analysis,
                            firm_value_prop, sender_name, sender_title, firm_name
                        )
                        email_drafts.append({"Company": company, "Draft": email})

                st.session_state.pe_advanced_outreach_results = email_drafts

        if 'pe_advanced_outreach_results' in st.session_state:
            st.success("✅ Email drafts generated successfully!")
            results = st.session_state.pe_advanced_outreach_results
            
            word_bytes = generate_word_document_from_drafts(results)
            st.download_button(
                "📥 Download All Drafts (.docx)", word_bytes, "advanced_outreach_drafts.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True
            )
            
            st.markdown("---")
            for item in results:
                st.subheader(item['Company'].replace('_', ' '))
                st.code(item['Draft'], language='text')
                st.markdown("---")
