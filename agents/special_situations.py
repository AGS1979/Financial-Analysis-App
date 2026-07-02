"""Agent Special Situations.

Generates M&A / spin-off / activist memos with an SOTP valuation module and an
infographic, using DeepSeek and FMP/yfinance data.
"""

import json
import pdfplumber
import re
import streamlit as st
import tempfile
import yfinance as yf

from config import require_env
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from typing import List, Dict, Tuple
from llm import llm
from utils.logging import log_audit_event, log_user_history, get_user_history
from utils.net import http_get


def special_situations_app():
    """
    Encapsulates the complete Agent Special Situations functionality,
    including memo generation, a valuation module, and infographic creation.
    """

    # ========== CONFIG & SETUP ==========
    # This section handles API key loading and global constants.
    _cfg = require_env("DEEPSEEK_API_KEY", "FMP_API_KEY")
    DEEPSEEK_API_KEY = _cfg["DEEPSEEK_API_KEY"]
    DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
    FMP_API_KEY = _cfg["FMP_API_KEY"]

    # ==========================
    # REPORT & INFOGRAPHIC STRUCTURES
    # ==========================
    REPORT_TEMPLATES = {
        "Spin-Off or Split-Up": """
Transaction Overview
ParentCo Post-Spin Outlook
SpinCo Investment Case
Valuation Analysis
Risks and Overhangs
""",
        "Mergers & Acquisitions": """
Deal Summary
Target Company Analysis
Buyer’s Rationale and Financing
Shareholder Vote & Antitrust Risk
Spread Analysis and Arbitrage Opportunity
""",
        "Bankruptcy / Distressed / Restructuring": """
Situation Summary
Capital Structure Analysis
Valuation and Recovery Scenarios
Reorganization Plan and Exit Timeline
Catalysts and Legal Risks
""",
        "Activist Campaign": """
Activist Background
Campaign Details
Company's Response and Governance Profile
Scenario Analysis
Valuation Impact
""",
        "Regulatory or Legal Catalyst": """
Legal/Regulatory Background
Outcome Scenarios
Financial and Strategic Implications
Market Reaction History
""",
        "Asset Sales or Carve-Outs": """
Transaction Overview
Strategic Impact
Use of Proceeds
Re-rating Potential
""",
        "Capital Raising or Buyback Catalyst": """
Transaction Mechanics
Capital Structure Post-Deal
Shareholder Implications
Buyback Analysis
"""
    }

    FALLBACK_META = [
        ("💼", "border-blue-600", "bg-blue-50"),
        ("🏢", "border-sky-600", "bg-sky-50"),
        ("🌐", "border-indigo-600", "bg-indigo-50"),
        ("🧩", "border-purple-600", "bg-purple-50"),
        ("📊", "border-green-600", "bg-green-50"),
        ("📈", "border-emerald-600", "bg-emerald-50"),
        ("👥", "border-yellow-600", "bg-yellow-50"),
        ("⚠️", "border-red-600", "bg-red-50"),
        ("💡", "border-pink-600", "bg-pink-50"),
        ("🧠", "border-gray-600", "bg-gray-50"),
    ]

    # ==========================
    # HELPER FUNCTIONS
    # ==========================

    # --- Text Extractors ---
    def extract_text_from_pdf(file):
        try:
            with pdfplumber.open(file) as pdf:
                return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
        except Exception as e:
            return f"[ERROR extracting PDF: {e}]"

    def extract_text_from_docx(file):
        try:
            doc = Document(file)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            return f"[ERROR extracting DOCX: {e}]"

    # --- Automated Financial Data Extraction for SOTP ---
    @st.cache_data(ttl=3600, show_spinner=False)
    def extract_financials_for_sotp(text: str, company_name: str) -> Dict:
        """Uses an LLM to extract structured financial data for SOTP analysis."""
        prompt = f"""
        Act as a financial analyst. From the documents about {company_name}'s spin-off, extract the key financials for each business segment.
        Identify the future Parent Company (likely the core remaining business, e.g., Tires), the Spin-Off company (e.g., Automotive/AUMOVIO), and any other major divisions mentioned that will be divested.

        Return a single, valid JSON object with the following structure. Use 'null' if a value is not found. All financial values should be in billions.
        {{
          "parent_co": {{
            "name": "Name of Parent Company (e.g., New Continental)",
            "segment_name": "Core Segment (e.g., Tires)",
            "sales": 14.0,
            "ebit": 1.9,
            "ebit_margin": 13.5
          }},
          "spin_co": {{
            "name": "Name of Spin-off (e.g., AUMOVIO)",
            "segment_name": "Spin-off Segment (e.g., Automotive)",
            "sales": 20.8,
            "ebit": 0.4,
            "ebit_margin": 1.9
          }},
          "other_divestitures": [
            {{
              "name": "Name of other segment to be sold (e.g., ContiTech)",
              "sales": 6.8,
              "ebit": 0.45,
              "ebit_margin": 6.7
            }}
          ],
          "currency": "EUR",
          "unit": "billion"
        }}

        CONTEXT:
        {text[:30000]}
        """
        try:
            content = llm.chat(
                [{"role": "user", "content": prompt}],
                temperature=0,
                response_format={"type": "json_object"},
                timeout=90,
            )
            response_json = json.loads(content)
            if 'parent_co' in response_json and 'spin_co' in response_json:
                return response_json
            return None
        except Exception as e:
            st.warning(f"Could not automatically extract structured financials: {e}")
            return None

    # --- NEWLY ADDED: Automated Multiple Assignment ---
    def get_valuation_multiples(sotp_financials: Dict) -> Dict:
        """Assigns reasonable valuation multiples based on business characteristics."""
        multiples = {}
        
        def get_multiple_for_segment(segment):
            if not segment or segment.get('sales') is None:
                return None
                
            margin = segment.get('ebit_margin')
            
            # Use EBIT multiple if profit is meaningful (e.g., > 100 million or >1% margin)
            if segment.get('ebit') and segment['ebit'] > 0.1:
                # Higher margin, stable businesses get higher multiples
                multiple_range = [7.0, 9.0] if margin and margin < 10.0 else [8.0, 10.0]
                return {"type": "EV/EBIT", "range": multiple_range, "metric": segment.get('ebit')}
            # Fallback to Sales multiple if profit is low, zero, or negative
            else:
                return {"type": "EV/Sales", "range": [0.4, 0.6], "metric": segment.get('sales')}

        if financials := sotp_financials.get('parent_co'):
            multiples['parent_co'] = get_multiple_for_segment(financials)
        if financials := sotp_financials.get('spin_co'):
            multiples['spin_co'] = get_multiple_for_segment(financials)
        
        for i, segment in enumerate(sotp_financials.get('other_divestitures', [])):
             multiples[f'other_divestitures_{i}'] = get_multiple_for_segment(segment)
             
        return multiples

    # --- Financial Data Fetchers ---
    @st.cache_data(ttl=3600, show_spinner=False)
    def resolve_company_to_ticker(company_name: str) -> str:
        prompt = f"What is the stock ticker (FMP-compatible) for the public company '{company_name}'?"
        try:
            ticker = llm.chat(
                [{"role": "user", "content": prompt}],
                temperature=0,
            ).strip()
            return re.sub(r'[^A-Z\.]', '', ticker)
        except:
            return None

    @st.cache_data(ttl=3600, show_spinner=False)
    def get_ev_ebitda_multiple(ticker: str, fmp_key: str) -> float:
        url = f"https://financialmodelingprep.com/api/v3/key-metrics-ttm/{ticker}?apikey={fmp_key}"
        try:
            r = http_get(url)
            data = r.json()
            if isinstance(data, list) and data:
                return float(data[0].get("enterpriseValueOverEBITDATTM", 0))
        except:
            return 0.0

    @st.cache_data(ttl=3600, show_spinner=False)
    def fetch_fundamentals_yf(ticker: str) -> Tuple[float, float, float]:
        """Returns (market_cap, net_debt, ttm_ebitda) via Yahoo Finance."""
        try:
            t = yf.Ticker(ticker)
            info = t.info or {}
            market_cap, total_debt, cash = info.get("marketCap", 0) or 0, info.get("totalDebt", 0) or 0, info.get("cash", 0) or 0
            net_debt = total_debt - cash
            ebitda = info.get("ebitda", 0) or 0
            return float(market_cap), float(net_debt), float(ebitda)
        except Exception:
            return 0.0, 0.0, 0.0

    # --- Text & Document Processors ---
    def clean_markdown(text):
        text = re.sub(r'^[ \t\-]{3,}$', '', text, flags=re.MULTILINE)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`{1,3}(.*?)`{1,3}', r'\1', text)
        text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'^- ', '• ', text, flags=re.MULTILINE)
        return text.strip()

    def truncate_safely(text, limit=30000):
        return text[:limit]

    def split_into_sections(raw_text: str, template: str) -> Dict[str, str]:
        sections = {}
        titles = [line.split('(')[0].strip() for line in template.strip().split('\n') if line.strip()]
        if not titles: return {"Investment Memo": raw_text.strip()}
        pattern = re.compile(r"^#+\s*(" + "|".join(map(re.escape, titles)) + r")\s*$", re.MULTILINE | re.IGNORECASE)
        matches = list(pattern.finditer(raw_text))
        if not matches:
            st.warning("Could not find structured headings in the AI's response.")
            return {"Investment Memo": raw_text.strip()}
        for i, match in enumerate(matches):
            title_from_text = match.group(1).strip()
            canonical_title = next((t for t in titles if t.lower() == title_from_text.lower()), title_from_text)
            start_of_content = match.end()
            end_of_content = matches[i + 1].start() if i + 1 < len(matches) else len(raw_text)
            content = raw_text[start_of_content:end_of_content].strip()
            if content: sections[canonical_title] = content
        return sections

    def format_memo_docx(memo_dict: dict, company_name: str, situation_type: str):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Aptos Display'
        style.font.size = Pt(11)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f"{company_name} – {situation_type} Investment Memo")
        title_run.font.name = 'Aptos Display'
        title_run.font.size = Pt(20)
        title_run.bold = True
        doc.add_paragraph()
        for section_title, content in memo_dict.items():
            heading = doc.add_paragraph()
            run = heading.add_run(section_title)
            run.bold = True
            run.font.size = Pt(14)
            heading.paragraph_format.space_after = Pt(6)
            for para in content.strip().split('\n\n'):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.3
        section = doc.sections[0]
        section.left_margin, section.right_margin, section.top_margin, section.bottom_margin = (Inches(0.75),)*4
        return doc

    # --- Core Memo Generator ---
    def generate_special_situation_note(
        company_name: str,
        situation_type: str,
        uploaded_files: list,
        valuation_mode: str = "Automated SOTP Analysis",
        parent_peers: str = "",
        spinco_peers: str = ""
    ):
        # 1) Extract text from uploaded documents
        combined_text = ""
        for file in uploaded_files:
            if file.name.endswith(".pdf"): combined_text += extract_text_from_pdf(file) + "\n"
            elif file.name.endswith(".docx"): combined_text += extract_text_from_docx(file) + "\n"
            else: combined_text += f"[Unsupported file: {file.name}]\n"

        # 2) Select the appropriate report structure
        structure = REPORT_TEMPLATES.get(situation_type)
        if not structure: raise ValueError(f"Unsupported situation type: {situation_type}")

        # 3) Build valuation section for Spin-Offs
        valuation_section = ""
        if situation_type == "Spin-Off or Split-Up":
            try:
                sotp_data = extract_financials_for_sotp(combined_text, company_name)
                if not sotp_data: raise ValueError("Financial data for SOTP not found in documents.")

                multiples = {}
                if valuation_mode == "Use Manual Peers":
                    st.info("Using manually provided peers for valuation.")
                    def process_peers(raw_peers):
                        names = [n.strip() for n in raw_peers.split(",") if n.strip()]
                        if not names: return None
                        tickers = [resolve_company_to_ticker(n) for n in names]
                        mults = [get_ev_ebitda_multiple(t, FMP_API_KEY) for t in tickers if t]
                        return round(sum(mults) / len(mults), 2) if mults else None
                    
                    parent_multiple = process_peers(parent_peers) or 8.0 # Fallback
                    spinco_multiple = process_peers(spinco_peers) or 8.0 # Fallback
                    
                    multiples['parent_co'] = {"type": "EV/EBIT", "range": [parent_multiple, parent_multiple], "metric": sotp_data.get('parent_co',{}).get('ebit')}
                    multiples['spin_co'] = {"type": "EV/EBIT", "range": [spinco_multiple, spinco_multiple], "metric": sotp_data.get('spin_co',{}).get('ebit')}
                else: # Default to Automated SOTP
                    st.info("Attempting automated SOTP analysis.")
                    multiples = get_valuation_multiples(sotp_data)

                results, total_low, total_high = [], 0, 0
                def calculate_segment_value(segment_key, segment_data):
                    mult_info = multiples.get(segment_key)
                    if not mult_info or not mult_info.get('metric') or mult_info.get('metric') <= 0: return None, 0, 0
                    low_val = mult_info['metric'] * mult_info['range'][0]; high_val = mult_info['metric'] * mult_info['range'][1]
                    results.append({
                        "name": segment_data.get('name', 'N/A'), "metric_val": mult_info['metric'], "multiple_type": mult_info['type'],
                        "multiple_range": f"{mult_info['range'][0]}x - {mult_info['range'][1]}x", "low_val": low_val, "high_val": high_val,
                        "unit": sotp_data.get('unit', 'billion'), "currency": sotp_data.get('currency', '$')})
                    return True, low_val, high_val
                
                _, low, high = calculate_segment_value('parent_co', sotp_data.get('parent_co'))
                total_low += low; total_high += high
                _, low, high = calculate_segment_value('spin_co', sotp_data.get('spin_co'))
                total_low += low; total_high += high
                for i, segment in enumerate(sotp_data.get('other_divestitures', [])):
                     _, low, high = calculate_segment_value(f'other_divestitures_{i}', segment)
                     total_low += low; total_high += high
                
                currency_symbol = '€' if sotp_data.get('currency') == 'EUR' else '$'
                unit_label = sotp_data.get('unit', 'billion')

                valuation_section += "## Valuation Analysis\n"
                valuation_section += f"Based on a Sum-of-the-Parts (SOTP) analysis, the implied total enterprise value is **{currency_symbol}{total_low:.1f} to {currency_symbol}{total_high:.1f} {unit_label}**.\n\n"
                valuation_section += "| Business Segment | Key Metric | Multiple | Implied Value Range |\n|:---|:---|:---|:---|\n"
                for res in results:
                    metric_str = f"{currency_symbol}{res['metric_val']:.1f} {unit_label} {res['multiple_type'].split('/')[1]}"
                    value_str = f"{currency_symbol}{res['low_val']:.1f} - {currency_symbol}{res['high_val']:.1f} {unit_label}"
                    valuation_section += f"| **{res['name']}** | {metric_str} | {res['multiple_range']} {res['multiple_type']} | {value_str} |\n"
                valuation_section += "\n[AI, please use this quantitative SOTP data to write a detailed narrative for the valuation analysis section.]"

            except Exception as e:
                st.warning(f"Automated SOTP analysis failed ({e}). Falling back to qualitative analysis.")
                valuation_section = ("## Valuation Analysis\n"
                                     "[AI, please generate a qualitative discussion on the potential valuation based on the documents.]")

        # 4) Assemble the corrected main prompt
        prompt = f"""You are an institutional investment analyst writing a professional memo on {company_name}'s {situation_type}.
        CONTEXT DOCUMENTS: \"\"\"{truncate_safely(combined_text)}\"\"\"
        VALUATION DATA (If available): {valuation_section}
        CRITICAL INSTRUCTIONS:
        1. Generate a detailed, well-written memo with comprehensive paragraphs for each section.
        2. Write in a narrative style. Each section should have at least 2-3 detailed paragraphs.
        3. You MUST use the exact section titles from the 'STRUCTURE' list below as level 2 markdown headings (e.g., `## Deal Summary`).
        4. If quantitative 'Valuation Analysis' data is provided, you MUST use it as the foundation for that section's narrative.
        5. **ABSOLUTELY NO** conversational introductions or conclusions (e.g., "Of course, here is..."). The response must start directly with the first markdown heading.

        STRUCTURE:
        {structure}"""

        # 5) Call LLM and process response
        raw_memo_text = llm.chat([{"role": "user", "content": prompt}], temperature=0.3)
        memo_dict_raw = split_into_sections(raw_memo_text, structure)
        memo_dict_cleaned = {title: clean_markdown(content) for title, content in memo_dict_raw.items()}
        
        # 6) Generate and save Word document
        doc = format_memo_docx(memo_dict_cleaned, company_name, situation_type)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            return tmp.name

    # --- Infographic Functions ---
    def extract_sections_from_docx_for_infographic(file, situation_type: str) -> Dict[str, str]:
        toc = REPORT_TEMPLATES.get(situation_type)
        if not toc: return {}
        expected_titles = {t.split('(')[0].strip().lower() for t in toc.strip().splitlines() if t.strip()}
        doc = Document(file)
        sections, current_heading, current_text = {}, None, []
        all_headings = [p.text.strip() for p in doc.paragraphs if p.runs and all(r.bold for r in p.runs if r.text.strip())]
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text: continue
            is_heading = text in all_headings and text.lower() in expected_titles
            if is_heading:
                if current_heading and current_text: sections[current_heading] = "\n".join(current_text).strip()
                current_heading, current_text = text, []
            elif current_heading: current_text.append(text)
        if current_heading and current_text: sections[current_heading] = "\n".join(current_text).strip()
        return sections

    def summarize_section_with_deepseek(section_title, section_text):
        prompt = f"""
You are an institutional research analyst preparing a financial infographic.
Your task is to summarize the provided section text into 3 to 5 concise bullet points.
Each point must be a single sentence, highlighting key insights clearly and professionally.
**CRITICAL INSTRUCTION:** Do NOT include any introductory or concluding phrases. Your response must begin directly with the first bullet point.
Section to Summarize:
\"\"\"{section_text}\"\"\"
"""
        return llm.chat([{"role": "user", "content": prompt}], temperature=0.3).strip()

    def build_infographic_html(company_name, sections):
        html = f"""
<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8" /><meta name="viewport" content="width=device-width, initial-scale=1.0"/><title>{company_name} – Infographic</title><script src="https://cdn.tailwindcss.com"></script><link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap" rel="stylesheet"><style>body {{ font-family: 'Inter', sans-serif; background-color: #f9fafb; color: #1f2937; }} .section-icon {{ font-size: 1.4rem; margin-right: 0.6rem; }}</style></head><body class="px-4 py-8 md:px-6 md:py-10 max-w-7xl mx-auto"><header class="text-center mb-12"><h1 class="text-3xl md:text-4xl font-bold text-gray-800 mb-2">{company_name} – Investment Memo Infographic</h1><p class="text-sm text-gray-500">Generated by ARANC'AI'</p></header><main class="grid grid-cols-1 md:grid-cols-2 lg:grid-cols-3 gap-6">"""
        with st.spinner("Summarizing sections for infographic..."):
            for idx, (title, section_text) in enumerate(sections.items()):
                icon, border_class, bg_class = FALLBACK_META[idx % len(FALLBACK_META)]
                try:
                    summary = summarize_section_with_deepseek(title, section_text)
                    cleaned_summary = summary.replace('**', '').replace('###', '').replace('##', '').replace('#', '')
                    lines = [line.lstrip("•*- ").strip() for line in cleaned_summary.split("\n") if line.strip()]
                    bullet_items = "\n".join(f"                    <li>{line}</li>" for line in lines)
                except Exception as e:
                    bullet_items = f"<li>Error generating summary: {e}</li>"
                    st.warning(f"Could not summarize section: '{title}'")
                html += f"""
        <div class="shadow-lg rounded-xl p-5 transition-transform hover:scale-[1.02] duration-300 ease-in-out border-l-4 {border_class} {bg_class}"><h2 class="text-lg font-semibold text-gray-800 mb-3 flex items-center"><span class="section-icon">{icon}</span>{title}</h2><ul class="list-disc text-sm text-gray-700 space-y-2 pl-5 leading-relaxed">{bullet_items}</ul></div>"""
        html += """
    </main><footer class="text-center mt-12"><p class="text-xs text-gray-400">This document is for informational purposes only and does not constitute investment advice.</p></footer></body></html>"""
        return html

    # ==========================
    # STREAMLIT UI & APP LOGIC
    # ==========================
    st.markdown("### 🔀 Agent Special Situations")
    st.subheader("Step 1: Generate Investment Memo")

    company_name_memo = st.text_input("Enter Company Name", key="company_name_memo")
    situation_type_memo = st.selectbox("Select Situation Type", options=list(REPORT_TEMPLATES.keys()), key="situation_type_memo")
    
    valuation_mode = "Automated SOTP Analysis"
    parent_peers_raw = ""
    spinco_peers_raw = ""

    if situation_type_memo == "Spin-Off or Split-Up":
        st.markdown("##### 🔍 Valuation Approach")
        valuation_mode = st.radio(
            "Choose a valuation approach for the SOTP analysis:",
            options=["Automated SOTP Analysis", "Use Manual Peers"],
            key="valuation_mode",
            horizontal=True,
            help="Automated analysis is the default. Choose manual to provide your own peer companies."
        )
        if valuation_mode == "Use Manual Peers":
            parent_peers_raw = st.text_area("Enter ParentCo Peer Company Names (comma-separated)", key="parent_peers_raw")
            spinco_peers_raw = st.text_area("Enter SpinCo Peer Company Names (comma-separated)", key="spinco_peers_raw")

    uploaded_files_memo = st.file_uploader("Upload Public Documents (PDF, DOCX)", accept_multiple_files=True, key="uploaded_files_memo")
    # --- NEW UI for Custom Prompt ---
    st.markdown("---")
    st.subheader("Advanced: Customize Memo Prompt")
    st.info("You can provide a custom prompt template below. The agent may require your prompt to ask for a memo based on a specific structure.")
    st.text_area(
        "Enter your custom prompt template:",
        placeholder="Enter your full custom prompt for the memo generation here...",
        height=250,
        key="situations_custom_prompt"
    )
    # --- END NEW UI ---
    if st.button("Generate Memo", type="primary"):
        if not company_name_memo or not situation_type_memo or not uploaded_files_memo:
            st.warning("Please fill in all fields and upload at least one document.")
        else:
            # --- ADD AUDIT LOG CALL ---
            log_audit_event(
                action_type="SITUATIONS_MEMO_GEN",
                status="STARTED",
                target_id=company_name_memo,
                details={"situation_type": situation_type_memo, "valuation_mode": valuation_mode, "files": [f.name for f in uploaded_files_memo]}
            )
            # ---
            with st.spinner("Generating memo... This may take a moment."):
                try:
                    memo_path = generate_special_situation_note(
                        company_name=company_name_memo,
                        situation_type=situation_type_memo,
                        uploaded_files=uploaded_files_memo,
                        valuation_mode=valuation_mode,
                        parent_peers=parent_peers_raw,
                        spinco_peers=spinco_peers_raw
                    )
                    st.session_state.memo_path = memo_path
                    st.session_state.company_name = company_name_memo
                    st.session_state.situation_type = situation_type_memo
                    
                    st.success("Memo generated successfully!")
                    # --- ADD AUDIT LOG CALL ---
                    log_audit_event(action_type="SITUATIONS_MEMO_GEN", status="SUCCESS", target_id=company_name_memo)
                    # ---
                    with open(memo_path, "rb") as f:
                        st.download_button(
                            label="Download Memo (.docx)",
                            data=f,
                            file_name=f"{company_name_memo}_{situation_type_memo}_Memo.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    # --- ADD AUDIT LOG CALL ---
                    log_audit_event(
                        action_type="SITUATIONS_MEMO_GEN",
                        status="FAILURE",
                        target_id=company_name_memo,
                        details={"error": str(e)}
                    )
                    # ---
                    st.error(f"An error occurred during memo generation: {e}")

    st.markdown("\n\n---\n\n")

    st.subheader("Step 2: Generate Infographic from Memo")
    st.info("After generating the memo, you can either upload it below or, if you just generated it, the app will use it automatically.")
    uploaded_memo_infographic = st.file_uploader("Upload the generated Memo (.docx)", type=["docx"], key="uploaded_memo_infographic")

    if st.button("Generate Infographic", type="primary"):
        memo_file_to_use = uploaded_memo_infographic
        if not memo_file_to_use and 'memo_path' in st.session_state:
            memo_file_to_use = st.session_state.memo_path
        
        company_name_infographic = st.session_state.get('company_name', '')
        situation_type_infographic = st.session_state.get('situation_type')

        if not memo_file_to_use or not company_name_infographic or not situation_type_infographic:
            st.warning("Please generate a memo first in Step 1, or upload a previously generated memo.")
        else:
            # --- ADD AUDIT LOG CALL ---
            log_audit_event(action_type="SITUATIONS_INFOGRAPHIC_GEN", status="STARTED", target_id=company_name_infographic)
            # ---
            with st.spinner("Extracting sections and generating infographic..."):
                try:
                    sections = extract_sections_from_docx_for_infographic(memo_file_to_use, situation_type_infographic)
                    if not sections:
                         st.error("Could not extract any sections from the document. Please ensure the memo was generated correctly with clear headings.")
                    else:
                        st.success(f"Successfully extracted {len(sections)} sections. Building infographic...")
                        html_content = build_infographic_html(company_name_infographic, sections)
                        
                        st.subheader("Infographic Preview")
                        st.components.v1.html(html_content, height=800, scrolling=True)
                        
                        # --- ADD AUDIT LOG CALL ---
                        log_audit_event(action_type="SITUATIONS_INFOGRAPHIC_GEN", status="SUCCESS", target_id=company_name_infographic)
                        # ---

                        st.download_button(
                            label="Download Infographic (.html)",
                            data=html_content,
                            file_name=f"{company_name_infographic}_Infographic.html",
                            mime="text/html"
                        )
                except Exception as e:
                    # --- ADD AUDIT LOG CALL ---
                    log_audit_event(
                        action_type="SITUATIONS_INFOGRAPHIC_GEN",
                        status="FAILURE",
                        target_id=company_name_infographic,
                        details={"error": str(e)}
                    )
                    # ---
                    st.error(f"An error occurred during infographic generation: {e}")
