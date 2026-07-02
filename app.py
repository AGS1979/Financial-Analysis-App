import uuid
from azure.ai.documentintelligence.models import ContentFormat, AnalyzeDocumentRequest
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from openai import AzureOpenAI # Use the Azure-specific client
from google.oauth2 import service_account
from google.cloud import aiplatform
from google.cloud import storage
from google.cloud import documentai_v1 as documentai
from google.cloud import dlp_v2
import vertexai
from vertexai.generative_models import GenerativeModel, Part
from st_supabase_connection import SupabaseConnection
import hashlib # For legacy SHA-256 password verification (migration only)
import hmac     # For constant-time legacy hash comparison
import bcrypt   # For password hashing
import html # Used to escape markdown characters
import io
from docx.enum.style import WD_STYLE_TYPE
from pinecone import Pinecone
import pickle
import streamlit as st
import os
import re
import fitz  # PyMuPDF
import faiss
import json
import requests
import markdown
import numpy as np
from datetime import datetime, timedelta, timezone
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from sentence_transformers import SentenceTransformer
from collections import defaultdict
import tempfile
import base64
from pathlib import Path
import streamlit.components.v1 as components
import pandas as pd
from io import BytesIO
from openai import OpenAI
import pdfplumber
import yfinance as yf
from typing import List, Dict, Tuple
from bs4 import BeautifulSoup
from utils import format_report_as_html
from PIL import Image, ImageDraw, ImageFont # Make sure PIL imports are at the top
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template
import toml
import config
from config import require_env, DEEPSEEK_API_URL
from utils.net import http_post, http_get
from utils.branding import get_base64_logo_image, load_logo
from utils.logging import log_audit_event, log_user_history, get_user_history
from auth.session import validate_session
from auth.ui import authentication_ui, whitelist_manager_ui
from agents.dcf import dcf_agent_app
from agents.special_situations import special_situations_app
from agents.esg import esg_analyzer_app
from agents.pe import pe_agent_app_azure
from agents.credit import agent_credit_app_azure
from agents.sentinel import agent_sentinel_app
from agents.ideagen import investment_pipeline_agent
from agents.real_time_sentinel import real_time_sentinel_app
from agents.commodity import commodity_forecasting_agent
from agents.risk_correlator import portfolio_risk_correlator_app
from agents.model_integrity import model_integrity_agent_app
from agents.tariff import tariff_impact_tracker_app


# --- Must be the first st.* command ---
st.set_page_config(
    page_title="ARANC'AI'",
    page_icon="📈",  # Adds a browser tab icon
    layout="wide"
)

# ==============================================================================
# APP BOOTSTRAP — stylesheet, config validation, and shared clients.
# Auth/config/util logic now lives in the auth/ and utils/ packages and config.py.
# ==============================================================================

def _inject_css():
    """Load the app stylesheet from static/styles.css once."""
    css_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "styles.css")
    with open(css_path, encoding="utf-8") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)


_inject_css()

# Validate the env vars required to boot, then build the shared clients (config.py).
config.validate_core_config()
openai_client = config.get_azure_client()
conn = config.get_conn()

# Re-export core secrets as module globals for the agent functions defined below.
DEEPSEEK_API_KEY = config.DEEPSEEK_API_KEY
FMP_API_KEY = config.FMP_API_KEY
OPENAI_API_KEY = config.OPENAI_API_KEY

# --- Page header ---
logo_base64 = get_base64_logo_image("logo.png")
st.markdown(
    f"""
    <div class="aranca-header">
        <div class="aranca-title">Welcome to ARANC'AI'</div>
        <div class="aranca-logo">
            <img src="data:image/png;base64,{logo_base64}" alt="Aranca Logo">
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

LOGO_OBJECT = load_logo()



# ==============================================================================
# 2. IPO INVESTMENT MEMO GENERATOR
# (Code from InvMemo.py and pipeline.py)
# ==============================================================================

# Add these imports to the top of your Python file
import streamlit as st
import requests
import re
import fitz  # PyMuPDF
import os
import tempfile
import numpy as np
import faiss
import markdown
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
from collections import defaultdict
from jinja2 import Template
from PyPDF2 import PdfReader
from sentence_transformers import SentenceTransformer
from pathlib import Path
from bs4 import BeautifulSoup # <-- ADD THIS IMPORT

def investment_memo_app():
    """
    Encapsulates the entire IPO Investment Memo Generator with Infographic and Q&A.
    This version is aligned with the advanced standalone module and supports both PDF and HTML.
    """
    
    # --- CONFIGURATION ---
    DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY")
    DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
    CHUNK_SIZE = 50

    if not DEEPSEEK_API_KEY:
        st.error("DeepSeek API key not found. Please add it to your Streamlit secrets.")
        return

    # --- HELPER FUNCTIONS (Memo & Pipeline) ---

    def clean_markdown(text):
        text = re.sub(r'#+\s*', '', text)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'_+', '', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'^[-*•]+\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'Section\s\d+[:.]?', '', text, flags=re.IGNORECASE)
        text = re.sub(r'(Next|Previous) section:.*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'This section .*?(focuses on|explores|explains).*?\.', '', text, flags=re.IGNORECASE)
        return text.strip()

    def extract_text_by_page(pdf_path):
        doc = fitz.open(pdf_path)
        return [page.get_text() for page in doc], len(doc)

    # MODIFICATION: New function to handle HTML files
    def extract_text_from_html(html_path):
        """Reads an HTML file and extracts all text content."""
        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'html.parser')
        
        for element in soup(["script", "style", "header", "footer", "nav"]):
            element.decompose() # Remove irrelevant tags
            
        text = soup.get_text(separator='\n', strip=True)
        # Consolidate whitespace
        cleaned_text = re.sub(r'\n{3,}', '\n\n', text)
        return cleaned_text

    def get_relevant_pages_chunked(text_by_page, user_query):
        total_pages = len(text_by_page)
        relevant_pages = set()
        for start in range(0, total_pages, CHUNK_SIZE):
            end = min(start + CHUNK_SIZE, total_pages)
            chunk_pages = text_by_page[start:end]
            prompt = (
                "Below are texts from a PDF. Identify only the page numbers (starting from 1) relevant to this query:\n"
                f"Query: {user_query}\n\n"
            )
            for i, text in enumerate(chunk_pages):
                snippet = text[:1000].replace('\n', ' ')
                prompt += f"\nPage {start + i + 1}: {snippet}\n"
            
            messages = [
                {"role": "system", "content": "You are an expert document analyst."},
                {"role": "user", "content": prompt}
            ]
            response = http_post(DEEPSEEK_API_URL, headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}, json={"model": "deepseek-chat", "messages": messages})
            response.raise_for_status()
            reply = response.json()['choices'][0]['message']['content']
            matches = re.findall(r'\d+', reply)
            for m in matches:
                if 1 <= int(m) <= total_pages:
                    relevant_pages.add(int(m))
        return sorted(relevant_pages)

    def extract_selected_pages_text(original_path, pages_to_keep):
        doc = fitz.open(original_path)
        return "\n".join(doc[p - 1].get_text() for p in pages_to_keep).strip()

    def extract_company_name(text):
        prompt = (
            "Extract only the legal name of the company from the following IPO or DRHP text. "
            "Return only the company name, nothing else.\n\n"
            f"{text[:10000]}"
        )
        messages = [
            {"role": "system", "content": "You are an expert in IPO documents."},
            {"role": "user", "content": prompt}
        ]
        response = http_post(DEEPSEEK_API_URL, headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}, json={"model": "deepseek-chat", "messages": messages})
        response.raise_for_status()
        
        # --- THE FIX IS HERE ---
        # Get the raw response and clean it by replacing newlines with spaces before returning.
        company_name = response.json()['choices'][0]['message']['content'].strip()
        return company_name.replace('\n', ' ')

    def find_relevant_text_for_section(full_text, section_title, keywords_map):
        """
        Searches the full document text for keywords related to a specific section title
        and returns a relevant chunk of text for the LLM to use as context.
        """
        # Get the specific keywords for the given section title
        keywords = keywords_map.get(section_title, [])
        if not keywords:
            # Fallback to a generic chunk if no keywords are defined for a title
            return full_text[:16000]

        # Create a regex pattern to find any of the keywords, ignoring case
        # This pattern looks for the keyword as a potential heading.
        pattern = re.compile(r'^\s*(' + '|'.join(map(re.escape, keywords)) + r')\s*$', re.IGNORECASE | re.MULTILINE)
        
        match = pattern.search(full_text)
        
        if match:
            # If a keyword is found, extract the text following it
            start_index = match.end()
            # Extract a substantial chunk of text to provide enough context
            context_text = full_text[start_index : start_index + 20000]
            return context_text
        else:
            # If no specific section is found, return the initial part of the document as a fallback.
            # This is not ideal but prevents errors. The key is a good keywords_map.
            st.warning(f"Could not find a specific section for '{section_title}'. Using initial text as fallback.")
            return full_text[:16000]

    def generate_memo_sections(filtered_text, custom_notes=""):
        section_titles = [
            "1. IPO Offer Details", "2. Company Overview", "3. Industry Overview and Outlook",
            "4. Business Model", "5. Financial Highlights",
            "6. Guidance and Outlook on future financial performance",
            "7. Peer Comparison and Competitors", "8. Risks", "9. Investment Highlights"
        ]

        # A map to find the right DRHP sections for each memo title. This is crucial.
        keywords_map = {
            "1. IPO Offer Details": ["Details of the Offer", "The Offer"],
            "2. Company Overview": ["Our Business", "Company Overview", "Summary of Business"],
            "3. Industry Overview and Outlook": ["Industry Overview"],
            "4. Business Model": ["Business Model", "Our Business"],
            "5. Financial Highlights": ["Financial Highlights", "Financial Information", "Restated Consolidated Financial Information", "Summary of Restated Consolidated Financial Information"],
            "6. Guidance and Outlook on future financial performance": ["Management's Discussion and Analysis", "Guidance and Outlook"],
            "7. Peer Comparison and Competitors": ["Peer Comparison", "Competitive Landscape", "Competitors"],
            "8. Risks": ["Risk Factors", "Risks"],
            "9. Investment Highlights": ["Investment Rationale", "Investment Highlights", "Our Competitive Strengths"]
        }
        
        sections = {}
        st.info("Generating memo sections with targeted context analysis...")
        progress_bar = st.progress(0)
        
        for i, title in enumerate(section_titles):
            st.write(f"Finding context for: **{title[3:]}**...")
            
            # --- THIS IS THE CORE FIX ---
            # Find relevant text specifically for THIS section, instead of using the same generic text for all.
            relevant_text = find_relevant_text_for_section(filtered_text, title, keywords_map)
            
            st.write(f"Generating content for: **{title[3:]}**...")
            
            prompt = (
                f"You are a professional financial analyst writing a Pre-IPO memo section titled: '{title[3:]}'.\n"
                "Based STRICTLY on the 'Relevant DRHP Text' provided below, generate approximately 500 words of clean, structured, analytical prose suitable for institutional investors.\n"
                "IMPORTANT RULES:\n"
                "- ONLY use information present in the provided text. Do not invent or hallucinate any data, figures, or facts.\n"
                "- If specific information (e.g., financial numbers, competitor names) is not in the text, state that the information is not available in the provided context.\n"
                "- Do not mention this is a memo. Do not start with the section title (e.g., avoid 'The Business Model is...').\n"
                "- Avoid phrases like 'In this section' or 'The provided text states...'. Write authoritatively.\n"
                "- Use plain, professional text only. Strictly avoid markdown (no asterisks, hashes, underscores).\n\n"
            )
            if custom_notes:
                prompt += f"USER'S CUSTOM FOCUS: {custom_notes.strip()}\n\n"
            
            prompt += f"Relevant DRHP Text:\n{relevant_text}" # We use the specifically found relevant_text
            
            messages = [
                {"role": "system", "content": "You are an expert financial analyst writing an investment memo based *only* on provided text."},
                {"role": "user", "content": prompt}
            ]
            
            try:
                response = http_post(
                    DEEPSEEK_API_URL, 
                    headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}, 
                    json={"model": "deepseek-chat", "messages": messages, "temperature": 0.2}
                )
                response.raise_for_status()
                raw_content = response.json()['choices'][0]['message']['content']
                
                # Clean the generated content
                cleaned = clean_markdown(raw_content)
                cleaned = re.sub(rf"^{re.escape(title[3:])}[\s:—-]*", "", cleaned, flags=re.IGNORECASE | re.MULTILINE)
                sections[title] = cleaned.strip()
                
            except requests.RequestException as e:
                st.error(f"API call failed for section '{title[3:]}': {e}")
                sections[title] = f"Error generating this section. {e}"

            progress_bar.progress((i + 1) / len(section_titles))
            
        st.success("All memo sections generated.")
        return sections

    def save_sections_to_word(sections_dict, company_name="Company", output_dir="documents"):
        os.makedirs(output_dir, exist_ok=True)
        
        # Sanitize the company name for the filename
        sanitized_company_name = re.sub(r'[\\/*?:"<>|]', "", company_name)
        sanitized_company_name = sanitized_company_name.replace(' ', '_')
        
        # Truncate to a reasonable length to avoid "File name too long" errors
        max_len = 50 
        if len(sanitized_company_name) > max_len:
            sanitized_company_name = sanitized_company_name[:max_len]

        filename = f"{sanitized_company_name}_PreIPO_Memo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        full_path = os.path.join(output_dir, filename)
        
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Aptos Display'
        style.font.size = Pt(11)

        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"{company_name} Pre-IPO Memo")
        title_run.font.name = 'Aptos Display'
        title_run.font.size = Pt(20)
        title_run.bold = True
        doc.add_paragraph()

        for title, body in sections_dict.items():
            heading = doc.add_paragraph()
            run = heading.add_run(title)
            run.bold = True
            run.font.name = 'Aptos Display'
            run.font.size = Pt(14)
            for para in body.strip().split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.add_paragraph()
        
        section = doc.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        
        doc.save(full_path)
        return full_path

    # MODIFICATION: The run_memo_pipeline is modified below in the main app flow
    # to handle the logic split between PDF and HTML.

    # --- HELPER FUNCTIONS (Infographic) ---

    def extract_raw_text_from_docx(docx_path):
        doc = Document(docx_path)
        return "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())

    def call_deepseek_summary(text, company_name):
        prompt = f"""
        You are an investment analyst. Summarize the key points for each section of the provided pre-IPO memo for {company_name}.
        For each section, provide 3-5 crisp bullet points. Each bullet point must be under 30 words.
        Format your response in markdown, with each section as a header.
        
        Memo Text:
        {text}
        """
        messages = [
            {"role": "system", "content": "You are a financial analyst specializing in concise summaries for infographics."},
            {"role": "user", "content": prompt}
        ]
        response = http_post(DEEPSEEK_API_URL, headers={"Authorization": f"Bearer {DEEPSEEK_API_KEY}"}, json={"model": "deepseek-chat", "messages": messages, "temperature": 0.3})
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    
    def parse_deepseek_response(summary_text):
        sections = defaultdict(list)
        current_section = None
        header_pattern = re.compile(r"^#+\s*(?:\d+\.\s*)?(.*?)\s*$", re.MULTILINE)
        
        parts = header_pattern.split(summary_text)
        if len(parts) > 1:
            for i in range(1, len(parts), 2):
                header = parts[i].strip()
                content = parts[i+1]
                bullets = re.findall(r'[-\*•]\s+(.*)', content)
                formatted_bullets = [re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', b.strip()) for b in bullets]
                if header and formatted_bullets:
                    sections[header] = formatted_bullets
        return dict(sections)

    def generate_infographic_html(docx_path, company_name, template_path="base_infographic.html"):
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"The template file '{template_path}' was not found.")
        raw_text = extract_raw_text_from_docx(docx_path)
        summary = call_deepseek_summary(raw_text, company_name)
        sections = parse_deepseek_response(summary)
        
        with open(template_path, "r", encoding="utf-8") as f:
            html_template = f.read()
            
        template = Template(html_template)
        return template.render(company_name=company_name, sections=sections)

    # --- HELPER FUNCTIONS (Q&A Engine) ---

    # MODIFICATION: Renamed and updated to handle both PDF and HTML
    class DocumentQueryEngine:
        def __init__(self, model_name="all-MiniLM-L6-v2"):
            self.api_key = DEEPSEEK_API_KEY
            self.embedder = SentenceTransformer(model_name)

        def _extract_chunks_from_pdf(self, path):
            reader = PdfReader(path)
            return [(i + 1, page.extract_text().strip()) for i, page in enumerate(reader.pages) if page.extract_text()]

        def _extract_chunks_from_html(self, path):
            full_text = extract_text_from_html(path)
            # Chunk HTML text by paragraphs (or a fixed character length if needed)
            paragraphs = [p.strip() for p in full_text.split('\n\n') if p.strip()]
            # Return with a chunk index instead of a page number
            return [(i + 1, para) for i, para in enumerate(paragraphs)]

        def answer_query(self, doc_path, query, top_k=3):
            # MODIFICATION: Select the correct chunking method based on file type
            if doc_path.endswith('.pdf'):
                chunks = self._extract_chunks_from_pdf(doc_path)
                source_label = "Page"
            elif doc_path.endswith('.html'):
                chunks = self._extract_chunks_from_html(doc_path)
                source_label = "Section" # Use a more generic label for HTML chunks
            else:
                raise ValueError("Unsupported file type for Q&A.")

            if not chunks: raise ValueError("No text could be extracted from the document.")
            
            source_ids, texts = zip(*chunks)
            text_embeddings = np.array(self.embedder.encode(texts, convert_to_numpy=True))
            
            index = faiss.IndexFlatL2(text_embeddings.shape[1])
            index.add(text_embeddings)
            
            query_embedding = self.embedder.encode([query])
            _, I = index.search(query_embedding, k=top_k)
            
            context_chunks = [(source_ids[i], texts[i]) for i in I[0]]
            
            messages = [{"role": "system", "content": "Answer the user's question based on the context provided from the document."}]
            for source_id, text in context_chunks:
                messages.append({"role": "user", "content": f"[Context from {source_label} {source_id}]:\n{text}"})
            messages.append({"role": "user", "content": f"Question: {query}"})
            
            response = http_post(DEEPSEEK_API_URL, headers={"Authorization": f"Bearer {self.api_key}"}, json={"model": "deepseek-chat", "messages": messages, "temperature": 0.2})
            response.raise_for_status()
            
            answer_md = response.json()["choices"][0]["message"]["content"]
            cited_sources = sorted([source_ids[i] for i in I[0]])
            return markdown.markdown(answer_md), cited_sources, source_label

    # Initialize session state
    if "memo_generated" not in st.session_state:
        st.session_state.memo_generated = False
    if "memo_path" not in st.session_state:
        st.session_state.memo_path = None
    if "doc_path" not in st.session_state: # MODIFICATION: Renamed for generality
        st.session_state.doc_path = None
    
    # --- Main App Flow ---
    st.markdown("### 📝 Agent Pre-IPO")
    st.markdown("Upload a DRHP or IPO prospectus to automatically generate a detailed investment memo, an infographic, and perform Q&A.")
    st.subheader("📤 1. Upload DRHP or IPO Prospectus")
    
    # MODIFICATION: Allow both PDF and HTML files
    uploaded_file = st.file_uploader(
        "Upload your PDF or HTML file",
        type=["pdf", "html"],
        key="prospectus_uploader"
    )
    
    if uploaded_file:
        # Save uploaded file to a temporary path with the correct extension
        suffix = Path(uploaded_file.name).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(uploaded_file.getbuffer())
            st.session_state.doc_path = tmp_file.name
            
        custom_focus = st.text_area(
            "Optional: Add custom notes to guide memo generation",
            key="memo_focus",
            help="Example: 'Focus on the competitive landscape in North America and risks related to supply chain.'"
        )

        # --- NEW UI for Custom Section Prompts ---
        st.markdown("---")
        st.subheader("Advanced: Customize Section Prompts")
        st.info("You can provide your own generation prompt for each memo section below. Leave a box empty to use the default.")
        
        section_titles = [
            "1. IPO Offer Details", "2. Company Overview", "3. Industry Overview and Outlook",
            "4. Business Model", "5. Financial Highlights",
            "6. Guidance and Outlook on future financial performance",
            "7. Peer Comparison and Competitors", "8. Risks", "9. Investment Highlights"
        ]

        for title in section_titles:
            st.markdown(f"##### Custom Prompt for: {title[3:]}")
            st.text_area(
                label=f"Custom prompt for {title[3:]}",
                placeholder=f"Enter your full custom prompt for the '{title[3:]}' section here...",
                height=150,
                key=f"memo_custom_prompt_{title.replace(' ', '_')}", # Unique key
                label_visibility="collapsed"
            )
        # --- END NEW UI ---

        # In your main app flow, under the "if st.button..."
        if st.button("📘 Generate Investment Memo", key="gen_memo"):
            # --- ADD AUDIT LOG CALL ---
            log_audit_event(
                action_type="PRE_IPO_MEMO_GEN",
                status="STARTED",
                target_id=uploaded_file.name,
                details={"custom_focus": custom_focus}
            )
            # ---
            with st.spinner("⏳ Analyzing document and generating memo... This may take a few minutes."):
                try:
                    full_text = ""
                    if st.session_state.doc_path.endswith(".pdf"):
                        # The extract_text_by_page function returns a list of strings (one per page)
                        text_by_page, _ = extract_text_by_page(st.session_state.doc_path)
                        # Join them all into one single large string
                        full_text = "\n".join(text_by_page)
                    
                    elif st.session_state.doc_path.endswith(".html"):
                        full_text = extract_text_from_html(st.session_state.doc_path)

                    if not full_text.strip():
                        raise ValueError("Could not extract text from the document.")
                    
                    company_name = extract_company_name(full_text)
                    st.session_state.pre_ipo_company_name = company_name # Save for later
                    
                    # --- CALL THE NEW, FIXED FUNCTION ---
                    # It now takes the full_text and handles context internally.
                    sections_dict = generate_memo_sections(full_text, custom_focus)
                    
                    memo_path = save_sections_to_word(sections_dict, company_name=company_name)

                    st.session_state.memo_generated = True
                    st.session_state.memo_path = memo_path
                    st.success("✅ Memo generated successfully!")
                    # --- ADD AUDIT LOG CALL ---
                    log_audit_event(action_type="PRE_IPO_MEMO_GEN", status="SUCCESS", target_id=company_name)
                    # ---
                    
                    # --- START: NEW HISTORY LOG CALL ---
                    log_user_history(
                        action_type="Pre-IPO Memo",
                        target_id=company_name,
                        summary=f"Generated Pre-IPO Memo for {company_name}",
                        details={"source_file": uploaded_file.name}
                    )
                    # --- END: NEW HISTORY LOG CALL ---


                except Exception as e:
                    # --- ADD AUDIT LOG CALL ---
                    log_audit_event(
                        action_type="PRE_IPO_MEMO_GEN",
                        status="FAILURE",
                        target_id=uploaded_file.name,
                        details={"error": str(e)}
                    )
                    # ---
                    st.error(f"❌ Error generating memo: {e}")
                    st.session_state.memo_generated = False
        
        # --- Post-Generation Options ---
        if st.session_state.memo_generated and st.session_state.memo_path:
            with open(st.session_state.memo_path, "rb") as f:
                st.download_button(
                    "📥 Download Memo (.docx)", f, Path(st.session_state.memo_path).name,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            st.markdown("---")
            st.subheader("🎨 2. Generate Infographic")
            if st.button("🖼️ Generate Infographic", key="gen_infographic"):
                company_name = st.session_state.get("pre_ipo_company_name", "Company")
                # --- ADD AUDIT LOG CALL ---
                log_audit_event(action_type="PRE_IPO_INFOGRAPHIC_GEN", status="STARTED", target_id=company_name)
                # ---
                with st.spinner("✨ Creating infographic summary..."):
                    try:
                        infographic_html = generate_infographic_html(st.session_state.memo_path, company_name)
                        
                        st.components.v1.html(infographic_html, width=1100, height=1000, scrolling=True)
                        
                        # --- ADD AUDIT LOG CALL ---
                        log_audit_event(action_type="PRE_IPO_INFOGRAPHIC_GEN", status="SUCCESS", target_id=company_name)
                        # ---
                                        
                        st.download_button(
                            label="📥 Download Infographic (.html)",
                            data=infographic_html,
                            file_name=f"{company_name.replace(' ', '_')}_Infographic.html",
                            mime="text/html"
                        )
                    except Exception as e:
                        # --- ADD AUDIT LOG CALL ---
                        log_audit_event(
                            action_type="PRE_IPO_INFOGRAPHIC_GEN",
                            status="FAILURE",
                            target_id=company_name,
                            details={"error": str(e)}
                        )
                        # ---
                        st.error(f"❌ Error generating infographic: {e}")

        # --- Q&A Section ---
        st.markdown("---")
        st.subheader("🔍 3. Ask Questions from the Document")
        query = st.text_input("Type your question (e.g., What are the key risk factors?)", key="memo_query")
        if query:
            # --- ADD AUDIT LOG CALL ---
            log_audit_event(
                action_type="PRE_IPO_QA",
                status="STARTED",
                target_id=st.session_state.doc_path,
                details={"query": query}
            )
            # ---
            with st.spinner("💬 Searching for answers in the document..."):
                try:
                    # MODIFICATION: Use the updated engine
                    engine = DocumentQueryEngine()
                    answer_html, cited_sources, source_label = engine.answer_query(st.session_state.doc_path, query)
                    st.markdown(answer_html, unsafe_allow_html=True)
                    st.caption(f"📄 Answer generated from information on {source_label.lower()}s: {', '.join(map(str, cited_sources))}")
                    # --- ADD AUDIT LOG CALL ---
                    log_audit_event(
                        action_type="PRE_IPO_QA",
                        status="SUCCESS",
                        target_id=st.session_state.doc_path,
                        details={"query": query}
                    )
                    # ---
                except Exception as e:
                    # --- ADD AUDIT LOG CALL ---
                    log_audit_event(
                        action_type="PRE_IPO_QA",
                        status="FAILURE",
                        target_id=st.session_state.doc_path,
                        details={"error": str(e)}
                    )
                    # ---
                    st.error(f"❌ Query Error: {e}")

# ==============================================================================
# 3. DCF AGENT
# (Code from app - DCFAgent.py and 1_📄_Report.py)
# ==============================================================================
                


# ==============================================================================
# 4. Agent Special Situations
# (Code from app - SpecialSituations.py)
# ==============================================================================

# ==============================================================================
# 5. ESG ANALYZER
# (Code from app-ESG.py and ESGComp.py)
# ==============================================================================




# ==============================================================================
# 6. Agent Portfolio (FINAL CORRECTED VERSION)
# ==============================================================================

def portfolio_agent_app(user_id: str):
    """
    A persistent agent to index and query company documents using Pinecone,
    with added capabilities for pre-defined, structured analysis.
    """
    import xml.etree.ElementTree as ET
    import json # Added for parsing risk assessment
    import html # Added for escaping HTML characters
    import pandas as pd # Added for reading Excel estimates
    
    # FMP API Key is still needed for transcripts and news
    FMP_API_KEY = os.environ.get("FMP_API_KEY")

    st.markdown("### 🗂️ Agent Portfolio")
    st.markdown("Upload company-specific documents for indexation.")

    # --- HELPER FUNCTIONS (Existing) ---
    def truncate_context(excerpts: list, max_chars: int = 120000) -> str:
        full_context = ""
        for excerpt in excerpts:
            if len(full_context) + len(excerpt) > max_chars:
                break
            full_context += excerpt
        return full_context

    def add_spacing_to_run_on_text(text: str) -> str:
        text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', text)
        text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', text)
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        return text

    def call_deepseek_model(prompt: str, is_json: bool = False) -> str:
        try:
            if not DEEPSEEK_API_KEY:
                st.error("DeepSeek API Key is not configured in secrets.")
                return "Error: API Key not available."
            
            headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}
            payload = {
                "model": "deepseek-chat", 
                "messages": [{"role": "user", "content": prompt}], 
                "temperature": 0.1, 
                "max_tokens": 8192
            }
            # Add response_format if JSON is expected
            if is_json:
                payload["response_format"] = {"type": "json_object"}

            response = http_post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=240)
            response.raise_for_status()
            
            raw_content = response.json()["choices"][0]["message"]["content"]
            # For non-JSON, clean up spacing. For JSON, return as is.
            return raw_content if is_json else add_spacing_to_run_on_text(raw_content)

        except Exception as e:
            st.error(f"An unexpected error occurred: {e}")
            return f"Error: {e}"


    def format_competitive_analysis_output(raw_text: str) -> str:
        """
        Robustly formats the Competitive Analysis output.
        Tries to parse as XML first, then falls back to heuristic HTML cleanup.
        """
        # --- Fallback Function for Messy HTML-like Text ---
        def fallback_formatter(text: str) -> str:
            # Remove the answer tags
            text = re.sub(r'</?answer>', '', text, flags=re.IGNORECASE)
            # Add headings for key sections
            text = re.sub(r'^\s*<competitive_landscape>\s*', '## Competitive Landscape\n', text, flags=re.IGNORECASE)
            text = re.sub(r'^\s*<opportunity_gaps>\s*', '## Opportunity Gaps\n', text, flags=re.IGNORECASE)
            text = re.sub(r'^\s*<prioritized_actions>\s*', '## Prioritized Actions\n', text, flags=re.IGNORECASE)
            text = re.sub(r'^\s*<sources>\s*', '## Sources\n', text, flags=re.IGNORECASE)
            # Format strong tags and paragraphs
            text = re.sub(r'<p><strong>(.*?)</strong>', r'\n### \1', text, flags=re.DOTALL)
            text = re.sub(r'<strong>(.*?)</strong>', r'**\1**', text, flags=re.DOTALL)
            text = re.sub(r'</?p>', '\n', text)
            # Format lists
            text = re.sub(r'<li>', '\n* ', text)
            # Clean up all remaining tags
            text = re.sub(r'<.*?>', '', text)
            # Normalize whitespace
            text = re.sub(r'\n\s*\n', '\n\n', text).strip()
            return text

        # --- Main Function Logic ---
        match = re.search(r'<answer>(.*)</answer>', raw_text, re.DOTALL)
        if not match:
            return fallback_formatter(raw_text) # Try fallback even if no answer tag
        
        content = match.group(1)
        
        try:
            # Try to parse as clean XML
            root = ET.fromstring(f"<root>{content}</root>") # Wrap in root for safety
            md_parts = []
            
            def get_text(element, path, default=''):
                node = element.find(path)
                return node.text.strip() if node is not None and node.text else default

            landscape = root.find('competitive_landscape')
            if landscape is not None:
                md_parts.append("## Competitive Landscape")
                for child in landscape:
                    if child.tag not in ['competitors', 'adjacent_disruptors']:
                        title = child.tag.replace('_', ' ').title()
                        md_parts.append(f"**{title}:** {child.text.strip() if child.text else ''}\n")
                md_parts.append("\n---\n## Key Competitors")

                if landscape.find('competitors') is not None:
                    md_parts.append("### Direct Competitors")
                    for comp in landscape.find('competitors').findall('competitor'):
                        name = get_text(comp, 'name')
                        pos = get_text(comp, 'positioning')
                        price = get_text(comp, 'pricing')
                        moves = get_text(comp, 'recent_strategic_moves')
                        md_parts.append(f"#### {name}\n**Positioning:** {pos}\n\n**Pricing:** {price}\n\n**Recent Moves:** {moves}\n")
                
                if landscape.find('adjacent_disruptors') is not None:
                    md_parts.append("### Adjacent-Space Disruptors")
                    for dis in landscape.find('adjacent_disruptors').findall('disruptor'):
                        name = get_text(dis, 'name')
                        pos = get_text(dis, 'positioning')
                        price = get_text(dis, 'pricing')
                        moves = get_text(dis, 'recent_strategic_moves')
                        md_parts.append(f"#### {name}\n**Positioning:** {pos}\n\n**Pricing:** {price}\n\n**Recent Moves:** {moves}\n")

            gaps = root.find('opportunity_gaps')
            if gaps is not None:
                md_parts.append("\n---\n## Identified Opportunity Gaps")
                if gaps.find('comparison') is not None:
                    md_parts.append(get_text(gaps, 'comparison') + '\n')
                
                levers = gaps.find('levers') if gaps.find('levers') is not None else gaps
                for i, lever in enumerate(levers, 1):
                    name = get_text(lever, 'name')
                    desc = get_text(lever, 'description')
                    exploit = get_text(lever, 'diageo_exploitation')
                    md_parts.append(f"### {i}. {name}\n{desc}")
                    if exploit: md_parts.append(f"\n**Current Exploitation:** {exploit}\n")

            actions = root.find('prioritized_actions')
            if actions is not None:
                md_parts.append("\n---\n## Prioritized Strategic Actions")
                md_parts.append("| Action | Impact (1-5) | Feasibility (1-5) | Rationale |")
                md_parts.append("| :--- | :---: | :---: | :--- |")
                for action in actions.findall('action'):
                    name = get_text(action, 'name', get_text(action, 'description'))
                    impact = get_text(action, 'impact', get_text(action, 'impact_score'))
                    feasibility = get_text(action, 'feasibility', get_text(action, 'feasibility_score'))
                    rationale = get_text(action, 'rationale')
                    md_parts.append(f"| **{name}** | {impact} | {feasibility} | {rationale} |")

            sources = root.find('sources')
            if sources is not None:
                md_parts.append("\n---\n## Sources")
                for source in sources.findall('source'):
                    md_parts.append(f"* {source.text.strip()}")
            
            return "\n".join(md_parts)
        except ET.ParseError:
            # If XML parsing fails, use the fallback formatter on the original text
            return fallback_formatter(raw_text)

    def parse_markdown_to_structure(markdown_text: str, analysis_type: str) -> list:
        if analysis_type == "Custom Query":
            return [("Custom Query Response", markdown_text)]
            
        structure = []
        heading_pattern = re.compile(r'^#+\s+.*$', re.MULTILINE)
        matches = list(heading_pattern.finditer(markdown_text))
        if not matches:
            if markdown_text.strip():
                structure.append(("Overview", markdown_text.strip()))
            return structure
        for i, match in enumerate(matches):
            heading_text = match.group(0).strip()
            start_of_content = match.end()
            end_of_content = matches[i + 1].start() if (i + 1) < len(matches) else len(markdown_text)
            content_text = markdown_text[start_of_content:end_of_content].strip()
            structure.append((heading_text, content_text))
        return structure

    def markdown_to_word_bytes(structured_data: list, company_name: str, analysis_type: str) -> bytes:
        doc = Document()
        styles = doc.styles
        def define_style(style_name, style_type, font_name, font_size, is_bold=False):
            try:
                style = styles[style_name]
            except KeyError:
                style = styles.add_style(style_name, style_type)
            font = style.font
            font.name = font_name
            font.size = Pt(font_size)
            font.bold = is_bold
            if style_type == WD_STYLE_TYPE.PARAGRAPH:
                style.base_style = styles['Normal']
            return style
        title_style = define_style('DocTitle', WD_STYLE_TYPE.PARAGRAPH, 'Aptos Display', 16, True)
        heading_style = define_style('DocHeading', WD_STYLE_TYPE.PARAGRAPH, 'Aptos Display', 12, True)
        body_style = define_style('DocBody', WD_STYLE_TYPE.PARAGRAPH, 'Aptos Display', 9, False)
        title_style.paragraph_format.space_after = Pt(18)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing = 1.15
        doc.add_paragraph(f"{analysis_type} - {company_name}", style=title_style)
        for heading, content in structured_data:
            cleaned_heading = re.sub(r'^#+\s*', '', heading).strip()
            doc.add_paragraph(cleaned_heading, style=heading_style)
            is_table = False
            lines = content.strip().split('\n')
            if len(lines) > 1 and '|' in lines[0] and re.match(r'^\s*\|?(:?-+:?\|)+(:?-+:?)?\s*$', lines[1]):
                is_table = True
            if is_table:
                headers = [h.strip() for h in lines[0].strip('|').split('|')]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = h
                for row_line in lines[2:]:
                    row_cells = table.add_row().cells
                    cells = [c.strip() for c in row_line.strip('|').split('|')]
                    for i, c in enumerate(cells):
                        if i < len(row_cells):
                            row_cells[i].text = c
            else:
                for para in content.split('\n'):
                    para = para.strip()
                    if not para: continue
                    if para.startswith(('* ', '- ')):
                        doc.add_paragraph(para[2:], style='List Bullet')
                    else:
                        para = para.replace('**', '')
                        doc.add_paragraph(para, style=body_style)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def format_analysis_as_html(markdown_text: str, title: str, sources: str) -> str:
        content_html = markdown.markdown(markdown_text, extensions=['tables'])
        # Logic to prevent repeating the word "Analysis"
        final_title = title if "analysis" in title.lower() else f"{title} Analysis"
        html_style = """
        <style>
            .analysis-container { font-family: 'Poppins', sans-serif; border: 1px solid #e0e0e0; border-radius: 8px; padding: 25px; background-color: #f9f9f9; margin-top: 20px; }
            .analysis-container h1, .analysis-container h2, .analysis-container h3, .analysis-container h4 { color: #00416A; }
            .analysis-container h2 {font-size: 1.5em; border-bottom: 2px solid #00416A; padding-bottom: 10px; margin-top: 0;}
            .analysis-container h3 {font-size: 1.2em; padding-bottom: 5px; margin-top: 25px; border-bottom: 1px solid #e6f1f6;}
            .analysis-container table { width: 100%; border-collapse: collapse; margin: 15px 0; }
            .analysis-container th, .analysis-container td { border: 1px solid #ddd; padding: 10px 14px; text-align: left; }
            .analysis-container th { background-color: #e6f1f6; font-weight: 600; }
            .analysis-container p { margin-bottom: 1em; line-height: 1.6; }
            .analysis-container ul, .analysis-container ol { padding-left: 1.5em; }
            .analysis-container li { margin-bottom: 0.75em; line-height: 1.6;}
            .analysis-container .sources { font-size: 0.85em; color: #555; margin-top: 25px; text-align: right; }
        </style>
        """
        return f"""
        {html_style}
        <div class="analysis-container">
            <h2>{final_title}</h2>
            {content_html}
            <div class="sources"><strong>Sources:</strong> {sources}</div>
        </div>
        """

    # --- HELPER FUNCTIONS FOR RISK ASSESSMENT ---
    def highlight_text(full_text: str, quote: str) -> str:
        """Finds a quote in a larger text and wraps it in a highlight tag."""
        # Escape special characters for regex and handle variations in whitespace
        safe_quote = re.escape(quote)
        pattern = re.compile(r'\s*'.join(safe_quote.split()), re.IGNORECASE)
        match = pattern.search(full_text)
        
        if match:
            start, end = match.span()
            # Get some context around the match
            context_start = max(0, start - 150)
            context_end = min(len(full_text), end + 150)
            context = full_text[context_start:context_end]
            
            # Highlight the matched phrase within the context
            highlighted_context = (
                html.escape(context[:start - context_start]) +
                f"<mark>{html.escape(context[start - context_start:end - context_start])}</mark>" +
                html.escape(context[end - context_start:])
            )
            return f"...{highlighted_context}..."
        return f"<i>(Could not locate exact quote for highlighting, but it was sourced from this document)</i><br>{html.escape(quote)}"

    def format_risk_assessment_html(risks_data: list, company_name: str, sources_str: str) -> str:
        """Creates a styled HTML report for the risk assessment."""
        styles = """
        <style>
            .risk-assessment-container { font-family: 'Poppins', sans-serif; background-color: #ffffff; padding: 20px; }
            .risk-card { border: 1px solid #e0e0e0; border-left: 5px solid #c0392b; border-radius: 8px; margin-bottom: 20px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }
            .risk-card-header { background-color: #f9f9f9; padding: 15px 20px; border-bottom: 1px solid #e0e0e0; }
            .risk-card-header h3 { margin: 0; font-size: 1.2em; color: #c0392b; }
            .risk-card-body { padding: 20px; }
            .risk-card-body h4 { font-size: 1.0em; color: #333; margin-top: 0; margin-bottom: 8px; text-transform: uppercase; letter-spacing: 0.5px;}
            .risk-card-body p { font-size: 0.95em; line-height: 1.6; color: #555; }
            .risk-source-snapshot { background-color: #fdf5f5; border: 1px dashed #e5b8b4; border-radius: 4px; padding: 15px; margin-top: 15px; font-family: monospace, monospace; font-size: 0.85em; line-height: 1.5; color: #444; }
            .risk-source-snapshot mark { background-color: #f6caca; padding: 2px 4px; border-radius: 3px; }
            .risk-sources-footer { font-size: 0.85em; color: #555; margin-top: 25px; text-align: right; }
        </style>
        """
        cards_html = ""
        for i, risk in enumerate(risks_data, 1):
            # +++ NEW LOGIC for snapshot display +++
            snapshot_content = ""
            if risk.get('snapshot_url'):
                snapshot_content = f'<img src="{risk["snapshot_url"]}" alt="Source Snapshot" style="width:100%; border-radius: 4px;">'
            else:
                snapshot_content = risk.get('highlighted_quote', 'Source text not available.')

            cards_html += f"""
            <div class="risk-card">
                <div class="risk-card-header"><h3>Risk #{i}: {html.escape(risk.get('risk_title', 'Untitled Risk'))}</h3></div>
                <div class="risk-card-body">
                    <h4>Summary</h4>
                    <p>{html.escape(risk.get('risk_summary', 'N/A'))}</p>
                    <h4>Potential Impact</h4>
                    <p>{html.escape(risk.get('potential_impact', 'N/A'))}</p>
                    <h4>Source Snapshot</h4>
                    <div class="risk-source-snapshot">{snapshot_content}</div>
                </div>
            </div>
            """
        
        return f"""
        {styles}
        <div class='risk-assessment-container'>
            <h2>Risk Assessment for {company_name}</h2>
            {cards_html}
            <div class="risk-sources-footer"><strong>Source Documents:</strong> {sources_str}</div>
        </div>
        """

    # --- MODIFIED: Helper function for Variant Perception ---
    def fetch_fmp_data(ticker: str) -> dict:
        """
        Fetches latest 3 transcripts and recent news from FMP.
        """
        if not FMP_API_KEY:
            st.error("FMP_API_KEY is not configured in secrets.")
            return {"error": "FMP_API_KEY is not configured in secrets."}

        data = {
            "transcripts": [], # Will be a list of dicts
            "news": [] # Will be a list of dicts
        }
        
        try:
            # 1. Fetch Latest Earnings Call Transcripts (up to 3)
            try:
                transcript_url = f"https://financialmodelingprep.com/api/v3/earning_call_transcript/{ticker}?apikey={FMP_API_KEY}"
                transcript_res = http_get(transcript_url, timeout=20)
                transcript_res.raise_for_status()
                transcript_data = transcript_res.json()
                
                # Get the most recent 3 transcripts
                if transcript_data and isinstance(transcript_data, list):
                    for item in transcript_data[:3]: # Limit to 3
                        data["transcripts"].append({
                            "content": item.get("content", "Transcript content not available."),
                            "quarter": item.get("quarter", "N/A"),
                            "year": item.get("year", "N/A")
                        })
            except Exception as e:
                st.warning(f"Failed to fetch FMP transcript for {ticker}: {e}")
                # Don't error out, just return empty list

            # 2. Fetch Recent News
            try:
                news_url = f"https://financialmodelingprep.com/api/v3/stock_news?tickers={ticker}&limit=20&apikey={FMP_API_KEY}"
                news_res = http_get(news_url, timeout=10)
                news_res.raise_for_status()
                news_data = news_res.json()
                if news_data and isinstance(news_data, list):
                    data["news"] = [{"title": item.get("title"), "text": item.get("text")} for item in news_data]
            except Exception as e:
                st.warning(f"Failed to fetch FMP news for {ticker}: {e}")
                # Don't error out

            return data

        except Exception as e:
            st.error(f"An unexpected error occurred in fetch_fmp_data: {e}")
            return {"error": str(e)}


    # --- Add this NEW helper function inside the main portfolio_agent_app function ---

    def create_and_upload_snapshot(supabase_client, namespace: str, company: str, source_file: str, page_number: int, quote: str) -> str:
        """
        Downloads a source PDF, creates a highlighted image of the relevant page,
        and uploads it to Supabase, returning the public URL.
        """
        import uuid

        # --- FIX 1: Add validation for missing or non-numeric page_number ---
        if page_number is None:
            st.warning(f"Snapshot skipped for '{source_file}': Page number was missing from metadata.")
            return None
        
        try:
            # Ensure page_number is an integer for library compatibility
            page_number = int(page_number)
            if page_number <= 0: raise ValueError("Page number must be positive.")
        except (ValueError, TypeError):
            st.warning(f"Snapshot skipped for '{source_file}': Invalid page number format ('{page_number}').")
            return None

        if not source_file.lower().endswith('.pdf'):
            return None

        source_bucket = "source-documents"
        snapshot_bucket = "risk_snapshots"
        safe_company_name = re.sub(r'[<>:"/\\|?*]', '_', company.strip())
        source_path = f"{namespace}/{safe_company_name}/{source_file}"
        
        try:
            file_bytes = supabase_client.storage.from_(source_bucket).download(path=source_path)
            
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                # --- FIX 2: Validate that the page number exists in the document ---
                if page_number > len(doc):
                    st.warning(f"Snapshot generation failed for {source_file}: Requested page {page_number}, but the document only has {len(doc)} pages.")
                    return None
                
                page = doc.load_page(page_number - 1)
                
                text_instances = page.search_for(quote, quads=True)
                if not text_instances:
                    words = quote.split()
                    if len(words) > 5:
                        short_quote = " ".join(words[:5])
                        text_instances = page.search_for(short_quote, quads=True)

                for inst in text_instances:
                    page.add_highlight_annot(inst)

                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                
                snapshot_filename = f"{uuid.uuid4()}.png"
                snapshot_path = f"{namespace}/{safe_company_name}/{snapshot_filename}"
                
                supabase_client.storage.from_(snapshot_bucket).upload(
                    path=snapshot_path,
                    file=img_bytes,
                    file_options={"content-type": "image/png", "upsert": "true"}
                )
                
                return supabase_client.storage.from_(snapshot_bucket).get_public_url(snapshot_path)

        except Exception as e:
            st.warning(f"Snapshot generation failed for {source_file} (Page {page_number}): {e}")
            return None

    # --- THIS IS THE ERRONEOUS BLOCK THAT HAS BEEN REMOVED ---
    
    @st.cache_resource
    def load_agent(user_id):
        import tiktoken
        class PortfolioAgent:
            def __init__(self, user_id: str, index_name: str = "portfolio-agent"):
                self.namespace = user_id
                pinecone_api_key = os.environ.get("PINECONE_API_KEY")
                if not pinecone_api_key:
                    st.error("Configuration error: PINECONE_API_KEY is missing or empty.")
                    raise RuntimeError("PINECONE_API_KEY is not configured.")
                try:
                    self.pc = Pinecone(api_key=pinecone_api_key)
                    if index_name not in [idx.name for idx in self.pc.list_indexes()]:
                        st.error(f"Pinecone index '{index_name}' was not found.")
                        raise NameError(f"Index '{index_name}' not found.")
                    self.index = self.pc.Index(index_name)
                    self.embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
                except Exception as e:
                    st.error(f"Failed to connect to Pinecone: {e}")
                    raise

            def _init_supabase(self):
                from supabase import create_client, Client
                # CORRECTED: Access secrets under the 'connections.supabase' structure
                url = os.environ.get("SUPABASE_URL")
                key = os.environ.get("SUPABASE_KEY")
                if not url or not key:
                    st.error("Supabase URL or Key is not configured in secrets for the agent. Please check [connections.supabase].")
                    return None
                return create_client(url, key)


            def sanitize_filename(self, name: str) -> str:
                return re.sub(r'[<>:"/\\|?*]', '_', name.strip())

            def _extract_text(self, file_content: bytes, filename: str) -> str:
                try:
                    if filename.lower().endswith(".pdf"):
                        with fitz.open(stream=file_content, filetype="pdf") as doc:
                            return "\n".join(page.get_text() for page in doc)
                    elif filename.lower().endswith(".docx"):
                        doc = Document(io.BytesIO(file_content))
                        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    elif filename.lower().endswith(".txt"):
                        return file_content.decode("utf-8")
                except Exception as e:
                    st.warning(f"Could not read {filename}: {e}")
                return ""

            def _chunk_text(self, file_content: bytes, filename: str, max_tokens: int = 512, overlap_tokens: int = 50) -> List[dict]:
                """Chunks text page by page and returns a list of dictionaries with text and page number."""
                if not filename.lower().endswith(".pdf"):
                    # For non-PDFs, treat as a single page
                    full_text = self._extract_text(file_content, filename)
                    # This part remains similar to before but for a single block
                    try:
                        enc = tiktoken.get_encoding("cl100k_base")
                    except Exception:
                        enc = tiktoken.get_encoding("gpt2")
                    tokens = enc.encode(full_text)
                    chunks = []
                    start = 0
                    while start < len(tokens):
                        end = start + max_tokens
                        chunk_text = enc.decode(tokens[start:end]).strip()
                        if chunk_text:
                            chunks.append({"text": chunk_text, "page_number": 1})
                        start += (max_tokens - overlap_tokens)
                    return chunks

                # PDF-specific page-by-page chunking
                chunks_with_pages = []
                try:
                    with fitz.open(stream=file_content, filetype="pdf") as doc:
                        try:
                            enc = tiktoken.get_encoding("cl100k_base")
                        except Exception:
                            enc = tiktoken.get_encoding("gpt2")
                        
                        for page_num, page in enumerate(doc, start=1):
                            page_text = page.get_text()
                            if not page_text.strip():
                                continue
                            
                            tokens = enc.encode(page_text)
                            start = 0
                            while start < len(tokens):
                                end = start + max_tokens
                                chunk_text = enc.decode(tokens[start:end]).strip()
                                if chunk_text:
                                    chunks_with_pages.append({"text": chunk_text, "page_number": page_num})
                                start += (max_tokens - overlap_tokens)
                except Exception as e:
                    st.warning(f"Error chunking PDF {filename}: {e}")

                return chunks_with_pages

            def _get_year_from_filename(self, filename: str) -> int:
                matches = re.findall(r'\b(20\d{2})\b', filename)
                return int(max(matches)) if matches else 0

            def add_documents(self, company: str, uploaded_files: list):
                safe_company_name = self.sanitize_filename(company)
                
                # Initialize Supabase client
                supabase_client = self._init_supabase()
                if not supabase_client:
                    return

                with st.status(f"Processing documents for {safe_company_name}...", expanded=True) as status:
                    total_vectors_processed = 0
                    for file in uploaded_files:
                        # 1. Upload original document to Supabase Storage
                        status.write(f"Uploading {file.name} to source document storage...")
                        file_bytes = file.getvalue()
                        supabase_path = f"{self.namespace}/{safe_company_name}/{file.name}"
                        
                        try:
                            # Use upsert=True to overwrite if it already exists
                            supabase_client.storage.from_("source-documents").upload(
                                path=supabase_path,
                                file=file_bytes,
                                file_options={"content-type": file.type, "upsert": "true"}
                            )
                        except Exception as e:
                            st.error(f"Failed to upload {file.name} to Supabase: {e}")
                            continue # Skip this file if upload fails

                        # 2. Clear old entries from Pinecone
                        status.write(f"Clearing old Pinecone entries for {file.name}...")
                        self.index.delete(
                            filter={"company": {"$eq": safe_company_name}, "source_file": {"$eq": file.name}},
                            namespace=self.namespace
                        )

                        # 3. Chunk text page by page
                        file_year = self._get_year_from_filename(file.name)
                        status.write(f"Chunking and embedding {file.name}...")
                        # Pass file_bytes to the new chunking function
                        chunks_with_pages = self._chunk_text(file_bytes, file.name)
                        
                        if not chunks_with_pages:
                            st.write(f"Skipping {file.name}: no text could be extracted or chunked.")
                            continue

                        # 4. Create vectors and upsert to Pinecone with page_number metadata
                        chunk_texts = [item['text'] for item in chunks_with_pages]
                        vectors = self.embedding_model.encode(chunk_texts).tolist()
                        
                        vectors_to_upsert = []
                        for i, item in enumerate(chunks_with_pages):
                            chunk_id = f"{safe_company_name}-{self.sanitize_filename(file.name)}-p{item['page_number']}-{i}"
                            metadata = {
                                "company": safe_company_name, 
                                "source_file": file.name, 
                                "original_text": item['text'], 
                                "year": file_year,
                                "page_number": item['page_number'] # <<< CRITICAL ADDITION
                            }
                            vectors_to_upsert.append({"id": chunk_id, "values": vectors[i], "metadata": metadata})

                        if not vectors_to_upsert:
                            continue

                        status.write(f"Upserting {len(vectors_to_upsert)} chunks to Pinecone...")
                        batch_size = 100
                        for i in range(0, len(vectors_to_upsert), batch_size):
                            self.index.upsert(vectors=vectors_to_upsert[i:i + batch_size], namespace=self.namespace)
                        total_vectors_processed += len(vectors_to_upsert)

                    if total_vectors_processed > 0:
                        st.success(f"Successfully indexed {total_vectors_processed} new document chunks for **{company}**.")
                    else:
                        st.warning("No new content was indexed.")

            def query(self, query_text: str, companies: List[str], k: int = 30) -> Tuple[str, str]:
                """
                Performs a query against the Pinecone index with an enhanced RAG pipeline
                for more accurate and synthesized financial analysis.
                """
                query_vector = self.embedding_model.encode(query_text).tolist()
                query_filter = {"company": {"$in": [self.sanitize_filename(c) for c in companies]}}

                # Increased k to 30 to provide a wider context for deeper synthesis
                results = self.index.query(
                    vector=query_vector,
                    top_k=k,
                    filter=query_filter,
                    include_metadata=True,
                    namespace=self.namespace
                )

                if not results.matches:
                    return "I could not find relevant information in the indexed documents.", ""

                # This sorting logic is correct and critical, it must be maintained.
                results.matches.sort(key=lambda m: m.metadata.get('page_number', 0))
                results.matches.sort(key=lambda m: m.metadata.get('year', 0), reverse=True)
                
                context_excerpts = [
                    f"Excerpt from '{m.metadata['source_file']}' (Year: {m.metadata.get('year', 'N/A')}, Page: {m.metadata.get('page_number', 'N/A')}):\n\"{m.metadata['original_text']}\"\n"
                    for m in results.matches
                ]
                
                source_docs = sorted(list(set(m.metadata['source_file'] for m in results.matches)))
                safe_context = truncate_context(context_excerpts)

                # --- THIS IS THE FINAL PROMPT, ENGINEERED FOR MAXIMUM DEPTH AND CONTEXT ---
                prompt = f"""You are a world-class senior equity research analyst from a top-tier investment bank, known for producing exceptionally detailed and insightful reports. Your task is to provide a comprehensive, multi-faceted answer to the user's question, synthesizing all available information into a robust, professional-grade analysis.

**Core Task & Rules:**
1.  **Direct Answer First:** Begin your response with a direct, one-sentence answer to the user's specific question. This should be the single most current and authoritative fact.
2.  **Construct a Detailed Narrative:** After the direct answer, you must build a detailed, multi-paragraph narrative that provides the deepest possible context. You are REQUIRED to search the context for and include the following types of information, if available:
    * **Specific Quantitative Drivers:** Go beyond generalities. Instead of "AI demand," extract specific drivers like "demand for 3nm and 5nm nodes" or "growth in the HPC platform."
    * **Forward-Looking Guidance:** If the user asks about full-year guidance, you MUST also find and include any specific guidance for the next quarter (e.g., "Q3 2025 revenue is guided to be between $X and $Y billion").
    * **Financial Nuances & Headwinds:** Find and include any commentary on factors affecting the headline numbers, such as "foreign exchange (FX) headwinds," "margin impact from new fabs," or "cost dilution."
    * **Management's Rationale:** Summarize management's qualitative commentary and the reasoning behind their projections.
3.  **The "First is Final" Rule for Specific Metrics:** The context is pre-sorted chronologically (most recent first). For any specific, directly comparable metric (e.g., FY2025 revenue growth %), you **MUST** use the value from the **EARLIEST** excerpt in the context. Ignore all subsequent, older values for that specific metric.
4.  **No Meta-Commentary:** **NEVER** mention "conflicting data" or "older guidance." Present a single, unified analysis based on the most current information.
5.  **Professional Formatting:** Structure your response with a clear summary, followed by detailed bullet points or short paragraphs to elaborate on the different aspects of the analysis (Drivers, Quarterly Outlook, etc.).

--- PRE-SORTED CONTEXT (Most Recent First) ---
{safe_context}
--- QUESTION ---
{query_text}
--- ANSWER ---
"""
                return call_deepseek_model(prompt), ", ".join(source_docs)

            def get_unindexed_analysis(self, analysis_type: str, company: str, context: str) -> str:
                """Generates analysis from provided text without querying Pinecone."""
                ANALYSIS_CONFIG = self._get_analysis_config()
                config = ANALYSIS_CONFIG.get(analysis_type)
                if not config:
                    return "Invalid analysis type selected."
                
                is_json = analysis_type == "Risk Assessment"
                system_prompt = config['system_prompt'].replace('{COMPANY_NAME}', company)
                prompt = f"{system_prompt}\n\nBase your analysis *only* on the following context:\n--- DOCUMENT CONTEXT ---\n{context}\n--- END CONTEXT ---"
                
                return call_deepseek_model(prompt, is_json=is_json)


            def _get_analysis_config(self):
                """Centralized configuration for all analysis types."""
                return {
                    # --- NEW: Variant Perception Config (MODIFIED) ---
                    "Variant Perception": {
                        "search_query": "N/A - Uses FMP API and User Upload",
                        "system_prompt": """Act as a 'variant perception' analyst. The current Wall Street consensus estimates (provided by the user) are as follows:

{CONSENSUS_ESTIMATES_TABLE}

You are also provided with a compilation of recent earnings transcripts (from user uploads and FMP API) and recent news headlines. Your task is to identify specific points that the market may be underappreciating or overly pessimistic about regarding these estimates. 

**Instructions:**
1.  Analyze the provided consensus estimates.
2.  Carefully read the compilation of transcripts. Pay close attention to subtle changes in management tone, forward-looking guidance, or new strategic initiatives that contradict or support a different outlook than the provided consensus. Note the source and date of each transcript.
3.  Review the recent news for any additional catalysts or risks.
4.  Frame your output as a potential mispricing thesis in a concise markdown report.

**Report Structure:**

## Consensus View
(Briefly summarize the consensus estimates provided.)

## Potential Variant Perception
(Detail your analysis here, using bullet points to highlight specific discrepancies, subtle management commentary, or underappreciated news that challenge the consensus. **Directly reference line items from the user's table where possible.**)

## Key Signposts to Watch
(List 2-3 key events or metrics that would validate or invalidate this variant perception.)

--- DATA ---

## 1. Consensus Estimates (from User)
{CONSENSUS_ESTIMATES_TABLE}

## 2. Recent Transcripts Compilation
{TRANSCRIPTS_COMPILATION}

## 3. Recent News
{NEWS_HEADLINES}
"""
                    },
                    "Quick Company Note": {
                        "search_query": "Comprehensive company profile including business overview, products, services, market position, key financial data like revenue, profit, margins, cash flow, EPS, balance sheet items (debt, cash), industry trends, competitive landscape, investment highlights, strengths, weaknesses, opportunities, threats, risk factors, and any red flags like impairments or governance issues.",
                        "system_prompt": """You are an expert equity research analyst from a top-tier investment bank. Your task is to generate a professional 'Quick Company Note' based ONLY on the provided document excerpts.
CRITICAL INSTRUCTION: The output MUST be in clean markdown format. Ensure there is always a single space between separate words, and between numbers and words. Do not concatenate words together.

Structure your response with the following headings:

# 1. Company Overview
(Provide a comprehensive summary of the company as a narrative text.)

# 2. Financial Performance
(Create a markdown table with the columns: 'Metric', 'Most Recent Fiscal Year', 'Prior Fiscal Year', and 'YoY Growth / Change'. Include key metrics like Revenue, Operating Profit, and Free Cash Flow. **IMPORTANT: If data for a specific year or metric is not available in the context, clearly mark that cell with 'N/A'**. Do not make up data. Present any financial figures you can find, even if the table is incomplete.)

# 3. Key Investment Highlights
(Generate a list of concise bullet points, starting each with `*`. For each bullet, **bold the key takeaway** at the beginning. Example: `* **Dominant Market Position:** The company holds the #1 spot...`)

# 4. Key Risks
(Generate a list of concise bullet points, starting each with `*`. For each bullet, **bold the primary risk factor**. Example: `* **Regulatory Headwinds:** The company faces potential...`)

# 5. Red Flags
(Generate a list of concise bullet points, starting each with `*`. Identify any potential red flags like impairments, governance issues, or high valuation uncertainty. **Bold the core issue** of each point.)

# 6. Analyst Commentary
(Write a short, concluding paragraph that synthesizes the key findings and provides a balanced view on the company's position.)
"""
                    },
                    "Competitive Analysis": {
                        "search_query": "High-level overview of the company, its industry, main products, and business strategy to provide initial context.",
                        "system_prompt": """<instructions> You are a top-tier strategy consultant with deep expertise in competitive analysis, growth loops, pricing, and unit-economics-driven product strategy. If information is unavailable, state that explicitly. </instructions>

<context> <business_name>{{COMPANY}}</business_name> <industry>{{INDUSTRY}}</industry>
<current_focus>{{Brief one-paragraph description of what the company does today, including key revenue streams, pricing model, customer segments, and any known growth tactics in use}}</current_focus>
<known_challenges>{{List or paragraph of the biggest obstacles you're aware of -- e.g., slowing user growth, rising CAC, regulatory pressure}}</known_challenges> </context>

<task> 1. Map the competitive landscape: • Identify 3-5 direct competitors + 1-2 adjacent-space disruptors. • Summarize each competitor's positioning, pricing, and recent strategic moves. 2. Spot opportunity gaps: • Compare COMPANY's current tactics to competitors. • Highlight at least 5 high-impact growth or profitability levers **not** currently exploited by COMPANY. 3. Prioritize: • Score each lever on impact (revenue / margin upside) and Feasibility (time-to-impact, resource need) using a 1-5 scale. • Recommend the top 3 actions with the strongest Impact x Feasibility. </task>

<approach> - Go VERY deep. Research far more than you normally would. Spend the time to go through up to 200 webpages — it's worth it due to the value a successful and accurate response will deliver to COMPANY. - Don't just look at articles, forums, etc. — anything is fair game... COMPANY/competitor websites, analytics platforms, etc. </approach>

<output_format> Return ONLY the following XML: <answer> <competitive_landscape> </competitive_landscape> <opportunity_gaps> </opportunity_gaps> <prioritized_actions> </prioritized_actions> <sources> </sources> </answer> </output_format>"""
                    },
                    "Cap Structure": {
                        "search_query": "Detailed information about the company's capital structure, including short-term and long-term debt instruments, maturity dates, coupon rates, leases, equity, and debt covenants.",
                        "system_prompt": """You are a senior credit analyst.
**CRITICAL INSTRUCTION: The entire output must be plain text. Do NOT use any asterisks (`*`) for formatting. Ensure proper spacing between all words and numbers.**
Based on the provided text, synthesize all information about the company's capital structure. Prioritize information from documents with the most recent year if there are conflicts.
Format the output with clear headings and narrative descriptions. Do NOT use markdown tables.
# Capital Structure Analysis
## Debt Instruments
(For each debt instrument, describe it in a sentence. e.g., "The company has 5.0% senior notes due 2028 with a principal of $500 million.")
## Key Ratios
(Describe any relevant ratios found in the text in a paragraph.)
## Covenants
(Describe any mentioned financial or operational covenants in a paragraph.)"""
                    },
                    "Debt Details": {
    "search_query": "Table of borrowings, bonds, notes, and other debt instruments. Short term borrowings, medium and long term borrowings. Unsecured bank loans, overdrafts, sustainability linked loans. US bonds, EMTN programme bonds. Principal, coupon, maturity date, covenants.",
    "system_prompt": """You are a senior credit analyst specializing in data extraction.
**CRITICAL INSTRUCTION: Your entire output must be in clean markdown format.**
Based on the provided text, your primary task is to find any detailed tables listing debt instruments (bonds, loans, etc.) and replicate them accurately. You must also summarize any surrounding text about covenants and maturity.

Format the output precisely as follows:

# Debt Details Analysis

## Debt Instruments
First, search the context for any tables that list borrowings, bonds, or other debt instruments. Recreate these tables in markdown format exactly as they appear, including all rows and columns you can identify. Pay close attention to columns like 'Instrument', 'Principal Amount', 'Maturity Date', 'Coupon/Rate', or similar.

## Key Covenants
After the table, under this heading, write a plain text paragraph summarizing any text that describes financial or procedural covenants.

## Maturity Profile
Under this heading, write a plain text paragraph summarizing any text that describes the debt maturity profile.
"""
},
                    "Litigations and Court Cases/Claims": {
                        "search_query": "Details on litigations, legal proceedings, lawsuits, court cases, regulatory investigations, and contingent liabilities.",
                        "system_prompt": """You are a legal analyst.
**CRITICAL INSTRUCTION: The entire output must be plain text. Do NOT use any asterisks (`*`) for formatting. Ensure proper spacing between all words and numbers.**
From the context provided, compile a report on all legal and regulatory matters. For each distinct case, write a paragraph detailing the nature of the claim, its status, and any potential financial impact. Prioritize information from documents with the most recent year if there are conflicts."""
                    },
                    # --- REPLACED "Investment Story" WITH "Risk Assessment" ---
                    "Risk Assessment": {
                        "search_query": "All mentions of risk factors, potential risks, challenges, threats, contingent liabilities, legal proceedings, and uncertainties.",
                        "system_prompt": """You are an expert risk analyst. Your task is to analyze the provided document excerpts for {COMPANY_NAME} in chronological order.

**CRITICAL INSTRUCTIONS:**
1.  Identify **4-5 of the most significant COMPANY-SPECIFIC risks**.
2.  You MUST differentiate between company-specific risks (e.g., reliance on a single supplier, a major lawsuit) and generic industry/market risks (e.g., economic downturns, general competition). **Focus exclusively on the company-specific ones.**
3.  For each identified risk, you must extract the **exact verbatim quote** from the source text that best describes it.
4.  Return a single, valid JSON object. The root key should be "risks", and its value should be a list of objects.
5.  Each object in the list must have these four keys:
    - "risk_title": A short, descriptive title for the risk (e.g., "Dependency on Key Personnel").
    - "risk_summary": A concise one-paragraph summary explaining the risk.
    - "potential_impact": A one-paragraph analysis of the potential financial or operational impact on the company.
    - "source_quote": The exact, verbatim quote from the document that describes the risk.

**EXAMPLE JSON OUTPUT STRUCTURE:**
{
  "risks": [
    {
      "risk_title": "Reliance on a Single Product Line",
      "risk_summary": "The company derives over 85% of its total revenue from the 'InnovateX' product, creating significant concentration risk. Any downturn in this product's market could severely affect financial performance.",
      "potential_impact": "A decline in InnovateX sales could lead to a sharp fall in revenue and profitability. It also makes the company vulnerable to new competitors targeting this specific niche. This could impact stock valuation and the ability to fund future R&D.",
      "source_quote": "Our InnovateX product line accounted for approximately 87% and 85% of our net product revenues in fiscal 2024 and 2023, respectively."
    }
  ]
}
"""
                    },
                    "Company Strategy": {
                        "search_query": "Information on corporate strategy, business objectives, future plans, growth initiatives, market expansion, product development, and strategic priorities.",
                        "system_prompt": """You are a strategy consultant.
**CRITICAL INSTRUCTION: The entire output must be plain text. Do NOT use any asterisks (`*`) for formatting. Ensure proper spacing between all words and numbers.**
Outline the company's core strategy. Use paragraphs for sections like 'Vision & Mission', 'Strategic Pillars', and 'Growth Initiatives'. Prioritize information from documents with the most recent year if there are conflicts."""
                    },
                    "Compare Investment Ideas": {
                        "search_query": "Comprehensive comparison of companies including business overview, financial performance (revenue, profit, margins), balance sheet strength (debt, cash, leverage), strategic initiatives, growth drivers, competitive landscape, market position, investment highlights, risk factors, and management outlook.",
                        "system_prompt": """You are a senior buy-side investment analyst tasked with producing a comparative analysis of several investment ideas.
**CRITICAL INSTRUCTION: The entire output must be plain text. Do NOT use any asterisks (`*`) for formatting. Ensure proper spacing between all words and numbers.**
Based ONLY on the provided document excerpts for the selected companies, generate a professional investment comparison memo. The analysis must be objective, data-driven, and written in flowing narrative paragraphs. Avoid using bullet points.

Structure your response with the following headings:
# Executive Summary
# Financial Performance Comparison
# Balance Sheet Health
# Strategic Outlook & Growth Drivers
# Risk Profile Comparison
# Analyst Recommendation
"""
                    },
                    "Management Meeting Prep": {
                        "search_query": "Recent CEO comments, guidance, outlook, transcripts, investor presentations, reports on business fundamentals including volumes, pricing, margins, cash flow, strategy, and confidence in public statements.",
                        "system_prompt": """You are an expert institutional public equity investor. Your task is to prepare a briefing document for a non-deal roadshow lunch with the CEO of {COMPANY_NAME}.

**GOAL:**
Your primary goal is to summarize the attached documents to help prepare for a 1-2 hour meeting. Your analysis must look for signals that indicate if the business and story are getting better, worse, or staying the same. This includes any indications that the core fundamentals of the business – volumes, pricing, margins & cash flow – are getting better or worse. You must identify very subtle clues. Pay special attention to comments from the CEO about guidance, outlook, and his level of confidence in those statements, including any language inflections relative to his last few public statements.

**BACKGROUND:**
The user is a dispassionate fact-finder trying to understand the trajectory of the company's fundamental metrics and whether the company is likely to be a long-term winner.

**KEY TOPICS:**
From the attached documents, pull a list of key topics that should be discussed.

**SOURCES:**
Please use primarily the documents provided and any documents directly from the company. Approach all company statements with an objective and skeptical lens. Be wary of blogs or biased sources. The analysis must be unbiased and fact-driven.

**RETURN FORMAT:**
1.  **Key Topics:** Start with a section outlining the key topics from the documents.
2.  **Key Questions for the CEO:** Provide the 3 key questions that can be asked in the meeting to inform whether the business & story are getting better, worse, or staying the same.
3.  **"Tells" to Listen For:** For each of the 3 questions, provide certain clues or "tells" to listen for in the CEO's response.
4.  **Broader Question List:** Provide a broader list of the most important 12-15 questions to ask.

**WARNINGS:**
Approach this analysis without bias. Remain completely objective and do not become influenced by any of these statements or anything that is biased to be bullish or bearish. The goal is to provide clues towards the future trajectory of the company's stock price, informed by the evolution of fundamentals and the narrative.
"""
                    }
                }

            # --- MODIFIED: get_predefined_analysis to accept user_estimates_table ---
            def get_predefined_analysis(
                self, 
                analysis_type: str, 
                companies: List[str], 
                k: int = 40, 
                user_estimates_table: str = None,
                transcripts_data: str = None,
                news_data: str = None
            ) -> Tuple[str, str, object]:
                
                ANALYSIS_CONFIG = self._get_analysis_config()
                config = ANALYSIS_CONFIG.get(analysis_type)
                if not config: return "Invalid analysis type selected.", "", None
                
                # --- NEW: Custom workflow for Variant Perception ---
                if analysis_type == "Variant Perception":
                    if not companies or len(companies) != 1:
                        return "Error: Please select exactly one company (ticker) for Variant Perception.", "", None
                    if not user_estimates_table:
                        return "Error: User estimates table was not provided.", "", None
                    if not transcripts_data:
                        return "Error: Transcripts data was not provided.", "", None
                    
                    ticker = self.sanitize_filename(companies[0])
                    template = config['system_prompt']

                    # Populate the prompt
                    final_prompt = template.replace('{COMPANY_NAME}', ticker)
                    final_prompt = final_prompt.replace('{CONSENSUS_ESTIMATES_TABLE}', user_estimates_table) 
                    final_prompt = final_prompt.replace('{TRANSCRIPTS_COMPILATION}', transcripts_data)
                    final_prompt = final_prompt.replace('{NEWS_HEADLINES}', news_data or "No recent news found.")
                    
                    response_text = call_deepseek_model(final_prompt, is_json=False)
                    # The 'sources' string is now built in the UI logic, so we return a placeholder
                    return response_text, "See UI for sources", None
                
                # --- Existing workflow for all other analysis types ---
                query_vector = self.embedding_model.encode(config["search_query"]).tolist()
                query_filter = {"company": {"$in": [self.sanitize_filename(c) for c in companies]}}
                
                # For Risk Assessment, sort by year
                sort_order = None
                if analysis_type == "Risk Assessment":
                        # This is a conceptual sort; Pinecone doesn't directly support sorting by metadata.
                        # We fetch more results and sort them client-side.
                        k = 60 # Fetch more to get a better chronological view

                results = self.index.query(vector=query_vector, top_k=k, filter=query_filter, include_metadata=True, namespace=self.namespace)

                if not results.matches: return f"Could not find any documents for this analysis.", "", None
                
                # Sort matches by year for chronological analysis in Risk Assessment
                if analysis_type == "Risk Assessment":
                    results.matches.sort(key=lambda m: m.metadata.get('year', 0))

                context_excerpts = [f"Excerpt from '{m.metadata['source_file']} (Year: {m.metadata.get('year', 'N/A')})':\n\"{m.metadata['original_text']}\"\n" for m in results.matches]
                source_docs = sorted(list(set(m.metadata['source_file'] for m in results.matches)))
                safe_context = truncate_context(context_excerpts)

                system_prompt = config['system_prompt']
                company_str = ', '.join(companies)

                if analysis_type == "Competitive Analysis":
                    system_prompt = system_prompt.replace('{{COMPANY}}', company_str)
                    system_prompt = system_prompt.replace('COMPANY', company_str)
                    system_prompt = system_prompt.replace('{{INDUSTRY}}', '(To be determined by your research)')
                    system_prompt = system_prompt.replace(
                        '{{Brief one-paragraph description of what the company does today, including key revenue streams, pricing model, customer segments, and any known growth tactics in use}}',
                        '(To be determined by your research based on the company name provided)'
                    )
                    system_prompt = system_prompt.replace(
                        '{{List or paragraph of the biggest obstacles you\'re aware of -- e.g., slowing user growth, rising CAC, regulatory pressure}}',
                        '(To be determined by your research based on the company name provided)'
                    )
                    prompt = f"{system_prompt}\n\nYou can use the following internal document context as a potential starting point, but your primary instruction is to perform the deep external research as detailed in the prompt above.\n--- DOCUMENT CONTEXT ---\n{safe_context}\n--- END CONTEXT ---\n\nProvide the analysis for '{company_str}'."
                else:
                    system_prompt = system_prompt.replace('{COMPANY_NAME}', company_str)
                    prompt = f"{system_prompt}\n\nBase your analysis *only* on the following context:\n--- DOCUMENT CONTEXT ---\n{safe_context}\n--- END CONTEXT ---\n\nProvide the analysis for '{company_str}'."
                
                # Check if JSON output is expected
                is_json_output = analysis_type == "Risk Assessment"
                response_text = call_deepseek_model(prompt, is_json=is_json_output)
                
                return response_text, ", ".join(source_docs), results.matches


            def get_indexed_companies(self) -> List[str]:
                all_companies = set()
                try:
                    # Query a small number of vectors just to get metadata
                    response = self.index.query(vector=[0.0]*384, top_k=1000, include_metadata=True, namespace=self.namespace)
                    for match in response.matches:
                        company = match.metadata.get("company")
                        if company:
                            all_companies.add(company)
                except Exception as e:
                    st.warning(f"Could not fetch indexed companies: {e}")
                return sorted(list(all_companies))

            def delete_company_data(self, company_name: str):
                safe_name = self.sanitize_filename(company_name)
                try:
                    self.index.delete(filter={"company": {"$eq": safe_name}}, namespace=self.namespace)
                    st.success(f"Successfully deleted all data for **{company_name}**.")
                except Exception as e:
                    st.error(f"Failed to delete data for {company_name}: {e}")
            
        try:
            return PortfolioAgent(user_id=user_id)
        except Exception as e:
            st.error(f"Failed to initialize Agent Portfolio: {e}")
            return None

    # --- Streamlit UI for the Agent Portfolio ---
    agent = load_agent(user_id=user_id)
    if not agent:
        st.stop()

    st.subheader("📁 Index New Company Documents")
    
    with st.form("indexing_form", clear_on_submit=True):
        new_company = st.text_input("Company Name", placeholder="e.g., RTX Corp, Microsoft")
        new_docs = st.file_uploader("Upload Documents (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
        if st.form_submit_button("Index Documents", type="primary"):
            if new_company and new_docs:
                agent.add_documents(new_company, new_docs)
                st.cache_resource.clear() # Clear cache to refresh company list
                st.rerun()
            else:
                st.warning("Please provide a company name and at least one document.")

    st.markdown("---")
    st.subheader("🔍 Analyze & Manage Companies")

    indexed_companies = agent.get_indexed_companies()
    if not indexed_companies:
        st.info("No companies have been indexed for your account yet. You can still use 'Management Meeting Prep' by uploading temporary documents.")

    st.markdown("#### Run Analysis")
    
    # --- UPDATED: analysis_options list ---
    analysis_options = [
        "Variant Perception", # <-- NEW
        "Quick Company Note", "Competitive Analysis", "Management Meeting Prep", 
        "Compare Investment Ideas", "Risk Assessment",
        "Cap Structure", "Debt Details", "Litigations and Court Cases/Claims",
        "Company Strategy", "Custom Query"
    ]
    analysis_choice = st.selectbox("Select Analysis Type", options=analysis_options)

    # --- NEW: Special UI for Management Meeting Prep ---
    if analysis_choice == "Management Meeting Prep":
        st.markdown("##### Management Meeting Preparation")
        with st.form("meeting_prep_form"):
            prep_company_name = st.text_input("Company Name for Analysis", help="The name of the company you are meeting with.")
            
            source_choice = st.radio(
                "Select Document Source",
                ("Use Indexed Documents", "Upload Temporary Documents"),
                horizontal=True, key="prep_source"
            )

            indexed_selection = None
            temp_docs = None

            if source_choice == "Use Indexed Documents":
                indexed_selection = st.multiselect("Select Indexed Company", options=indexed_companies)
            else:
                temp_docs = st.file_uploader(
                    "Upload Documents for this analysis only (will not be indexed)",
                    type=["pdf", "docx", "txt"], accept_multiple_files=True
                )

            # --- NEW UI for Custom Prompt ---
            st.markdown("---")
            st.subheader("Advanced: Customize Meeting Prep Prompt")
            st.text_area(
                "Enter your custom prompt for the analysis:",
                placeholder="Enter your full custom prompt for the 'Management Meeting Prep' analysis here...",
                height=250,
                key="meeting_prep_custom_prompt"
            )
            # --- END NEW UI ---

            submitted = st.form_submit_button("🚀 Generate Meeting Prep", use_container_width=True)

            if submitted:
                if not prep_company_name:
                    st.error("Please enter the Company Name for the analysis.")
                elif source_choice == "Use Indexed Documents" and not indexed_selection:
                    st.error("Please select at least one indexed company.")
                elif source_choice == "Upload Temporary Documents" and not temp_docs:
                    st.error("Please upload at least one temporary document.")
                else:
                    with st.spinner(f"Preparing briefing for {prep_company_name}..."):
                        analysis_md, sources, pinecone_matches = "", "", None
                        
                        if source_choice == "Use Indexed Documents":
                            analysis_md, sources, _ = agent.get_predefined_analysis(
                                analysis_choice, companies=indexed_selection
                            )
                        else: # Upload Temporary Documents
                            context_list = []
                            source_names = []
                            for doc in temp_docs:
                                text = agent._extract_text(doc.getvalue(), doc.name)
                                if text:
                                    context_list.append(text)
                                    source_names.append(doc.name)
                            full_context = "\n\n---\n\n".join(context_list)
                            analysis_md = agent.get_unindexed_analysis(
                                analysis_choice, prep_company_name, full_context
                            )
                            sources = ", ".join(source_names)
                        
                        if "Error:" in analysis_md or "Could not find" in analysis_md or not analysis_md.strip():
                            st.error(analysis_md or "Failed to generate a response from the model.")
                        else:
                            structured_report = parse_markdown_to_structure(analysis_md, analysis_choice)
                            report_html = format_analysis_as_html(analysis_md, analysis_choice, sources)
                            word_bytes = markdown_to_word_bytes(structured_report, prep_company_name, analysis_choice)
                            
                            st.session_state['analysis_output'] = {
                                "html": report_html,
                                "word": word_bytes,
                                "company_name": prep_company_name,
                                "analysis_type": analysis_choice
                            }

    # --- MODIFIED: New UI flow for Variant Perception ---
    elif analysis_choice == "Variant Perception":
        st.info("ℹ️ This analysis requires a ticker, user-uploaded estimates, and transcripts.")
        
        vp_ticker = st.text_input("Enter Company Ticker (FMP Compatible)", placeholder="e.g., AAPL, MSFT")
        vp_estimates_file = st.file_uploader(
            "1. Upload Street Estimates (.xlsx)", 
            type=["xlsx"], 
            help="File must have a 'line items' column and columns for 'FY25', 'FY26', etc."
        )
        vp_transcript_files = st.file_uploader(
            "2. Upload Transcripts (Optional - will be combined with FMP data)", 
            type=["pdf", "docx", "txt"], 
            accept_multiple_files=True,
            help="Upload transcripts if FMP has no data or if you have more recent files."
        )

        if st.button("🚀 Run Variant Perception", use_container_width=True, type="primary"):
            if not vp_ticker or not vp_estimates_file:
                st.error("Please provide both a company ticker and an estimates file.")
                st.stop()

            with st.spinner(f"Running Variant Perception for {vp_ticker}..."):
                # 1. Process Estimates File
                user_estimates_table_str = None
                try:
                    df = pd.read_excel(vp_estimates_file)
                    if 'line items' not in [str(col).lower() for col in df.columns]:
                        st.error("Upload failed: Estimates Excel file must contain a 'line items' column.")
                        st.stop()
                    user_estimates_table_str = df.to_markdown(index=False)
                except Exception as e:
                    st.error(f"Failed to read Excel file: {e}")
                    st.stop()

                transcripts_compilation = []
                sources_list = ["User-Uploaded Estimates"]

                # 2. Process User-Uploaded Transcripts (Priority)
                if vp_transcript_files:
                    sources_list.append("User-Uploaded Transcripts")
                    for doc in vp_transcript_files:
                        text = agent._extract_text(doc.getvalue(), doc.name)
                        if text:
                            transcripts_compilation.append(f"--- TRANSCRIPT (Source: User Upload '{doc.name}') ---\n{text}\n\n")

                # 3. Process FMP Data (Transcripts & News)
                fmp_data = fetch_fmp_data(vp_ticker)
                if "error" in fmp_data:
                    st.error(f"Error fetching FMP data: {fmp_data['error']}")
                    st.stop()

                fmp_transcripts = fmp_data.get("transcripts", [])
                fmp_news = fmp_data.get("news", [])

                if fmp_transcripts:
                    sources_list.append("FMP API (Transcripts)")
                    for item in fmp_transcripts: # Already limited to 3 in fetch function
                        transcripts_compilation.append(f"--- TRANSCRIPT (Source: FMP API - Q{item.get('quarter')} {item.get('year')}) ---\n{item.get('content', 'N/A')}\n\n")

                # 4. Check for Failure
                if not transcripts_compilation:
                    st.error("No transcripts found. FMP had no data for this ticker, and no transcripts were uploaded. Please upload at least one transcript to run this analysis.")
                    st.stop()

                # 5. Compile News
                news_str = ""
                if fmp_news:
                    sources_list.append("FMP API (News)")
                    news_list = [
                        f"* **{item.get('title', 'No Title')}**: {item.get('text', 'No content')}" 
                        for item in fmp_news
                    ]
                    news_str = "\n".join(news_list)

                # 6. Call Agent
                final_transcripts_str = "\n".join(transcripts_compilation)
                analysis_md, _, _ = agent.get_predefined_analysis(
                    analysis_type=analysis_choice,
                    companies=[vp_ticker], # Pass ticker as the "company"
                    user_estimates_table=user_estimates_table_str,
                    transcripts_data=final_transcripts_str,
                    news_data=news_str
                )
                
                sources_str = ", ".join(sources_list)

                # 7. Process and Save Output
                if "Error:" in analysis_md or not analysis_md.strip():
                    st.error(analysis_md or "Failed to generate a response from the model.")
                else:
                    structured_report = parse_markdown_to_structure(analysis_md, analysis_choice)
                    report_html = format_analysis_as_html(analysis_md, analysis_choice, sources_str)
                    word_bytes = markdown_to_word_bytes(structured_report, vp_ticker, analysis_choice)
                    
                    st.session_state['analysis_output'] = {
                        "html": report_html,
                        "word": word_bytes,
                        "company_name": vp_ticker,
                        "analysis_type": analysis_choice
                    }

    # --- Existing UI for ALL OTHER analysis types ---
    else:
        if not indexed_companies:
             st.info("Please index documents for a company to run this analysis type.")
        else:
            selected_companies = st.multiselect("Select Company/Companies to Analyze", options=indexed_companies, default=indexed_companies[0] if indexed_companies else [])
            
            user_query = ""
            if analysis_choice == "Custom Query":
                user_query = st.text_area("Ask a question about the selected companies' documents")

            # --- NEW UI for Custom Prompt ---
            # Show custom prompt UI for all non-custom, non-variant-perception analyses
            if analysis_choice not in ["Custom Query", "Variant Perception"]:
                st.markdown("---")
                st.subheader("Advanced: Customize Analysis Prompt")
                if analysis_choice == "Risk Assessment":
                    st.warning("Your custom prompt must request a specific JSON output for the report to generate correctly.")
                st.text_area(
                    "Enter your custom prompt for the analysis:",
                    placeholder=f"Enter your full custom prompt for the '{analysis_choice}' analysis here...",
                    height=250,
                    key="portfolio_custom_prompt"
                )
            # --- END NEW UI ---


            if st.button("🚀 Run Analysis", use_container_width=True):
                proceed = False
                if not selected_companies:
                    st.warning("Please select at least one company.")
                elif analysis_choice == "Compare Investment Ideas" and len(selected_companies) < 2:
                    st.warning("Please select at least two companies for comparison.")
                elif analysis_choice == "Custom Query" and not user_query.strip():
                    st.warning("Please enter a question for the custom query.")
                else:
                    proceed = True

                if proceed:
                    with st.spinner(f"Running '{analysis_choice}' analysis for {', '.join(selected_companies)}..."):
                        analysis_md, sources, pinecone_matches = "", "", None
                        if analysis_choice == "Custom Query":
                            analysis_md, sources = agent.query(user_query, selected_companies)
                        else:
                            # This now handles all *other* predefined analyses
                            analysis_md, sources, pinecone_matches = agent.get_predefined_analysis(
                                analysis_type=analysis_choice, 
                                companies=selected_companies
                            )
                        
                        company_name_for_doc = selected_companies[0] if len(selected_companies) == 1 else "Multiple Companies"
                        
                        if "Error:" in analysis_md or "Could not find" in analysis_md or not analysis_md.strip():
                            st.error(analysis_md or "Failed to generate a response from the model.")
                        else:
                            report_html = ""
                            word_bytes = b""
                            
                            # --- Custom workflow for Risk Assessment ---
                            if analysis_choice == "Risk Assessment":
                                try:
                                    from thefuzz import fuzz
                                    supabase_client = agent._init_supabase()
                                    data = json.loads(analysis_md)
                                    risks = data.get("risks", [])
                                    
                                    for risk in risks:
                                        quote = risk.get("source_quote", "")
                                        risk['snapshot_url'] = None

                                        if not quote or not pinecone_matches or not supabase_client:
                                            risk['highlighted_quote'] = "Source text not available."
                                            continue

                                        best_match_meta = None
                                        highest_score = 0
                                        
                                        for match in pinecone_matches:
                                            score = fuzz.partial_ratio(quote.lower(), match.metadata['original_text'].lower())
                                            if score > highest_score:
                                                highest_score = score
                                                best_match_meta = match.metadata
                                        
                                        MIN_MATCH_SCORE = 75
                                        if highest_score >= MIN_MATCH_SCORE and best_match_meta:
                                            snapshot_url = create_and_upload_snapshot(
                                                supabase_client=supabase_client,
                                                namespace=user_id,
                                                company=company_name_for_doc,
                                                source_file=best_match_meta.get('source_file'),
                                                page_number=best_match_meta.get('page_number'),
                                                quote=quote
                                            )
                                            risk['snapshot_url'] = snapshot_url
                                        else:
                                            risk['highlighted_quote'] = (
                                                f"<i>(Could not find a high-confidence match for the quote in source documents.)</i><br>"
                                                f"<b>LLM-Generated Quote:</b> {html.escape(quote)}"
                                            )

                                    report_html = format_risk_assessment_html(risks, company_name_for_doc, sources)
                                    word_bytes = b"" # No Word doc for this type
                                    st.session_state['analysis_output'] = {
                                        "html": report_html,
                                        "word": word_bytes,
                                        "company_name": company_name_for_doc,
                                        "analysis_type": analysis_choice
                                    }

                                except json.JSONDecodeError:
                                    st.error("Failed to parse the Risk Assessment response from the AI. The format was not valid JSON.")
                                    st.text_area("Raw Response:", analysis_md, height=200)

                            else: # --- Existing workflow for all other types ---
                                if analysis_choice == "Competitive Analysis":
                                    analysis_md = format_competitive_analysis_output(analysis_md)

                                structured_report = parse_markdown_to_structure(analysis_md, analysis_choice)
                                report_html = format_analysis_as_html(analysis_md, analysis_choice, sources)
                                word_bytes = markdown_to_word_bytes(structured_report, company_name_for_doc, analysis_choice)
                                
                                if report_html:
                                    st.session_state['analysis_output'] = {
                                        "html": report_html,
                                        "word": word_bytes,
                                        "company_name": company_name_for_doc,
                                        "analysis_type": analysis_choice
                                    }

    # --- MODIFIED: Shared Output Display Area ---
    if 'analysis_output' in st.session_state:
        output = st.session_state.pop('analysis_output') # Get and remove to prevent re-display on rerun
        
        # Display a generic success message for all analysis types.
        st.success("✅ Analysis complete. Your report is ready for download below.")
        
        # This section no longer displays any HTML output in the UI.
        
        d1, d2 = st.columns(2)
        safe_filename = re.sub(r'[\s/]', '_', output["analysis_type"])
        doc_name = output["company_name"]
        
        # --- MODIFICATION: Don't show Word download for Risk Assessment ---
        if output["analysis_type"] != "Risk Assessment":
            d1.download_button(
                label="📥 Download as Word (.docx)",
                data=output["word"],
                file_name=f"{safe_filename}_{doc_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            d1.empty() # Keep layout consistent

        d2.download_button(
            label="📥 Download as HTML (.html)",
            data=output["html"].encode("utf-8"),
            file_name=f"{safe_filename}_{doc_name}.html",
            mime="text/html",
            use_container_width=True
        )


    if indexed_companies:
        st.markdown("---")
        st.markdown("#### Manage Data")
        
        company_to_delete = st.selectbox("Select Company to Delete", options=[""] + indexed_companies, key="delete_select")
        if st.button("🗑️ Delete All Data for This Company", type="secondary"):
            if company_to_delete:
                agent.delete_company_data(company_to_delete)
                st.cache_resource.clear()
                st.rerun()
            else:
                st.warning("Please select a company to delete.")



    

# ==============================================================================
# 7. TARIFF IMPACT TRACKER (NEW MODULE)
# ==============================================================================




# ==============================================================================
# 8. Agent PE (Azure POWERED)
# ==============================================================================



# ==============================================================================
# 9. Agent Credit (Azure POWERED) - NEW
# ==============================================================================




# ==============================================================================
# 10. Model Integrity Agent (Azure POWERED) - NEW
# ==============================================================================




# ==============================================================================
# 11. Agent Sentinel (Proactive Monitoring) - NEW
# ==============================================================================




# ==============================================================================
# 12. Agent Sentinel (Proactive Monitoring) - NEW
# ==============================================================================



# ==============================================================================
# 13. Real-Time Risk & Compliance Sentinel (Workflow)
# ==============================================================================




# ==============================================================================
# 14. Commodity Price Forecasting Agent (FINAL VERSION V2)
# ==============================================================================

# (Commodity Forecaster imports moved with the agent to agents/commodity.py)




# ==============================================================================
# 15. Portfolio Risk Correlator (NEW AGENT)
# ==============================================================================





# ==============================================================================
# 16. HISTORY PAGE FUNCTION (NEW)
# ==============================================================================

def show_history_page():
    """
    Renders a full page for the user's recent activity, using the app's CSS.
    """
    st.markdown("### 📜 Your Recent Activity")
    st.markdown("Here is a log of your most recent actions on the platform.")
    st.markdown("---")

    # Fetch a larger number of items since it's a full page
    history_items = get_user_history(limit=25) 

    if not history_items:
        st.info("You have no recent activity.")
    else:
        html_items = []
        for item in history_items:
            # Format the timestamp
            timestamp = pd.to_datetime(item['created_at']).strftime('%b %d, %Y at %I:%M %p')
            
            tags_html = ""
            
            # Action Tag
            if item.get('action_type'):
                safe_action = html.escape(item["action_type"])
                tags_html += f'<div class="history-tag" title="{safe_action}"><span class="history-tag-label">Action:</span> {safe_action}</div>'
            
            # Target Tag
            if item.get('target_id'):
                safe_target = html.escape(item["target_id"])
                tags_html += f'<div class="history-tag" title="{safe_target}"><span class="history-tag-label">Target:</span> {safe_target}</div>'
            
            # Parameters Tags
            if item.get('details'):
                clean_details = {k: v for k, v in item['details'].items() if v is not None}
                for key, value in clean_details.items():
                    clean_key = html.escape(key.replace('_', ' ').title())
                    value_str = html.escape(str(value))
                    tags_html += f'<div class="history-tag" title="{value_str}"><span class="history-tag-label">{clean_key}:</span> {value_str}</div>'

            # Assemble the full HTML for this one card
            safe_summary = html.escape(item['summary'])
            html_items.append(f"""
            <div class="history-item">
                <div class="history-summary">{safe_summary}</div>
                <div class="history-timestamp">Performed on {timestamp}</div>
                <div class="history-tags">
                    {tags_html}
                </div>
            </div>
            """)
        
        # Join all cards into a single HTML string
        full_history_html = "\n".join(html_items)
        
        # Render the HTML directly to the page. The page will scroll naturally.
        st.markdown(full_history_html, unsafe_allow_html=True)

# ==============================================================================
# 16. MAIN APP ROUTER (CORRECTED AND COMPLETE)
# ==============================================================================

def main():
    """
    Main function to run the Streamlit app with authentication and routing.
    This version includes a corrected height calculation for the agent card component.
    """
    
    if not authentication_ui():
        st.stop()
    
    if not validate_session():
        st.warning("Your session has expired or another session has been started with your credentials.")
        st.info("Please log in again.")
        for key in ['logged_in', 'username', 'session_token']:
            if key in st.session_state:
                del st.session_state[key]
        st.button("Reload")
        st.stop()

    if 'st_supabase_connection' not in st.session_state:
        try:
            supabase_url = os.environ.get("SUPABASE_URL")
            supabase_key = os.environ.get("SUPABASE_KEY")
            st.session_state.st_supabase_connection = SupabaseConnection('supabase', supabase_url=supabase_url, supabase_key=supabase_key)
        except Exception as e:
            st.error(f"Error initializing Supabase connection: {e}")
            st.stop()

    # --- DYNAMIC AGENT VISIBILITY LOGIC ---
    
    # Load the config file
    with open("config.toml", "r") as f:
        config = toml.load(f)

    # Get the permissions from the loaded config
    permissions = config.get("user_permissions", {})
    current_user = st.session_state.get("username")
    visible_agents = permissions.get("__DEFAULT__", ["🏠 Welcome"]) 
    if current_user in permissions:
        visible_agents = permissions[current_user]

    # --- SIDEBAR DEFINITION ---
    with st.sidebar:
        st.title("ARANC'AI'")
        st.write(f"Welcome, **{st.session_state.username}**")
        st.markdown("---")
        app_mode = st.radio(
            "Choose a tool:",
            options=visible_agents,
            key="app_tool_choice"
        )
        st.markdown("---")
        if st.button("Logout"):
            conn.client.table("users").update({"active_session_token": None}).eq("email", st.session_state['username']).execute()
            for key in st.session_state.keys():
                del st.session_state[key]
            st.rerun()
        st.info("App powered by Aranca.")

    st.markdown("---")

    # --- ROUTER LOGIC ---
    if app_mode == "Real-Time Sentinel":
        real_time_sentinel_app(user_id=st.session_state.username, client=openai_client)
    elif app_mode == "Agent IdeaGen": 
        investment_pipeline_agent()
    elif app_mode == "Portfolio Risk Correlator":
        portfolio_risk_correlator_app(client=openai_client)
    elif app_mode == "Agent Credit":
        agent_credit_app_azure()
    elif app_mode == "Model Integrity Agent":
        model_integrity_agent_app()
    elif app_mode == "Agent Sentinel":
        agent_sentinel_app()
    elif app_mode == "Agent PE":
        pe_agent_app_azure()
    elif app_mode == "Agent Pre-IPO":
        investment_memo_app()
    elif app_mode == "DCF Ginny":
        dcf_agent_app(client=openai_client, FMP_API_KEY=FMP_API_KEY)
    elif app_mode == "Agent Special Situations":
        special_situations_app()
    elif app_mode == "ESG Analyzer":
        esg_analyzer_app()
    elif app_mode == "Agent Portfolio":
        portfolio_agent_app(user_id=st.session_state.username)
    elif app_mode == "Tariff Impact Tracker":
        tariff_impact_tracker_app(DEEPSEEK_API_KEY=DEEPSEEK_API_KEY, FMP_API_KEY=FMP_API_KEY, logo_base64_string=logo_base64)
    elif app_mode == "Commodity Forecaster":
        commodity_forecasting_agent(client=openai_client)
    elif app_mode == "History":
        show_history_page()
    else: # This is the "🏠 Welcome" page
        st.markdown('<p class="welcome-subtitle">A unified platform for advanced financial analysis.</p>', unsafe_allow_html=True)
        st.info("👈 **Select an agent from the sidebar to begin.**")
        
        # --- The history section has been removed ---
        
        st.subheader("Available Agents")

        # Define all possible agent cards in a master list
        ALL_AGENT_DETAILS = [
            {"name": "History", "title": "📜 History", "description": "View a full log of your recent activity and generated analyses."},
            {"name": "Agent IdeaGen", "title": "💡 Agent IdeaGen", "description": "Discover new investment ideas by screening the market based on a specific theme or set of custom criteria."},
            {"name": "Agent PE", "title": "🔒 Agent PE", "description": "Analyze confidential IMs and teasers with enterprise-grade secured environment."},
            {"name": "DCF Ginny", "title": "📈 DCF Ginny", "description": "Generate a document-driven Discounted Cash Flow (DCF) analysis using public data or your own financials."},
            {"name": "ESG Analyzer", "title": "🌍 ESG Analyzer", "description": "Extract and compare key ESG metrics from sustainability reports to benchmark corporate performance."},
            {"name": "Agent Pre-IPO", "title": "📝 Agent Pre-IPO", "description": "Upload a DRHP/IPO PDF to automatically generate a detailed investment memo and perform Q&A."},
            {"name": "Agent Credit", "title": "🔒 Agent Credit", "description": "Analyze confidential credit agreements, indentures, and loan documents in a secure environment."},
            {"name": "Agent Portfolio", "title": "🗂️ Agent Portfolio", "description": "Index company-specific documents (10-Ks, earnings calls) and perform Q&A across your entire portfolio."},
            {"name": "Tariff Impact Tracker", "title": "📈 Tariff Impact Tracker", "description": "Analyze earnings calls or filings to extract mentions of tariffs and their financial impact."},
            {"name": "Agent Special Situations", "title": "📊 Agent Special Situations", "description": "Analyze events like M&A, spin-offs, and activist campaigns by uploading relevant documents to generate a summary memo."},
            {"name": "Agent Sentinel", "title": "📡 Agent Sentinel", "description": "Proactively monitor portfolio companies for key news, filings, and events."},
            {"name": "Model Integrity Agent", "title": "🛡️ Model Integrity Agent", "description": "Audit Excel financial models for errors, hard-codes, and inconsistencies."},
            {"name": "Commodity Forecaster", "title": "🌾 Commodity Forecaster", "description": "Forecast commodity prices using time-series data, technical indicators, and news sentiment analysis."},
            {"name": "Real-Time Sentinel", "title": "🚨 Real-Time Sentinel", "description": "Provides a real-time warning system for compliance issues and tail risks."},
            {"name": "Portfolio Risk Correlator", "title": "🧬 Portfolio Risk Correlator", "description": "Upload documents for multiple companies to identify and visualize hidden, correlated risks across your portfolio."}
        ]

        # Filter the cards based on the user's permissions
        visible_agent_details = [agent for agent in ALL_AGENT_DETAILS if agent["name"] in visible_agents]

        # Generate the HTML for the visible cards
        card_html_list = []
        for agent in visible_agent_details:
            card_html_list.append(f"""
            <div class="agent-card" title="{agent['description']}">
                <div class="agent-title">{agent['title']}</div>
                <div class="agent-description">{agent['description']}</div>
            </div>
            """)

        # Combine the CSS and the dynamic HTML cards into a single block
        full_html = f"""
        <html>
            <head>
                <style>
                .agent-grid {{
                    display: flex;
                    flex-wrap: wrap;
                    gap: 20px;
                }}
                .agent-card {{
                    flex: 1 1 30%;
                    min-width: 300px;
                    display: flex;
                    flex-direction: column;
                    background-color: #f8f9fa;
                    border: 1px solid #e0e0e0;
                    border-radius: 8px;
                    padding: 20px;
                    transition: box-shadow 0.2s ease-in-out;
                    font-family: 'Poppins', sans-serif;
                }}
                .agent-card:hover {{
                    box-shadow: 0 4px 12px rgba(0,0,0,0.1);
                }}
                .agent-title {{
                    font-size: 1.1rem;
                    font-weight: 600;
                    color: #1e1e1e;
                    margin-bottom: 10px;
                }}
                .agent-description {{
                    font-size: 0.95rem;
                    color: #4a4a4a;
                    line-height: 1.5;
                }}
                </style>
            </head>            <body>
                <div class="agent-grid">
                    {''.join(card_html_list)}
                </div>
            </body>
        </html>
        """
        
        num_rows = (len(visible_agent_details) + 2) // 3
        height_px = num_rows * 220
        
        components.html(full_html, height=height_px)

if __name__ == "__main__":
    main()